
"""
Script FINAL BERSIH - ganti gambar berdasarkan posisi tepat yang diketahui.
"""
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

SRC = r'c:\Users\ASUS\Documents\Codex\Losarijaya2\TA2_14694_REVISI_v2.docx'
DST = r'c:\Users\ASUS\Documents\Codex\Losarijaya2\TA2_14694_REVISI_v2.docx'

ARTIFACT_DIR = r'C:\Users\ASUS\.gemini\antigravity-ide\brain\8d8f779d-6960-4d6b-8d92-000b7d2c5da9'
files = os.listdir(ARTIFACT_DIR)

IMG_USECASE   = os.path.join(ARTIFACT_DIR, [f for f in files if 'usecase_diagram' in f][0])
IMG_ACT_ADMIN = os.path.join(ARTIFACT_DIR, [f for f in files if 'activity_diagram_admin' in f][0])
IMG_ACT_KASIR = os.path.join(ARTIFACT_DIR, [f for f in files if 'activity_diagram_kasir' in f][0])
IMG_CLASS     = os.path.join(ARTIFACT_DIR, [f for f in files if 'class_diagram_sistem' in f][0])
IMG_ERD       = os.path.join(ARTIFACT_DIR, [f for f in files if 'erd_diagram_sistem' in f][0])

NS_WD = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
NS_WP_I = '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline'
NS_WP_A = '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}anchor'

def has_image(p):
    return bool(p._element.findall(f'.//{NS_WP_I}') or p._element.findall(f'.//{NS_WP_A}'))

def clear_and_set_image(doc, para, img_path, width_inches):
    """Bersihkan gambar lama dari paragraf lalu set gambar baru."""
    p_elem = para._element
    for run_elem in list(p_elem.findall(f'{NS_WD}r')):
        if run_elem.findall(f'.//{NS_WP_I}') or run_elem.findall(f'.//{NS_WP_A}'):
            p_elem.remove(run_elem)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.add_picture(img_path, width=Inches(width_inches))
    print(f"  -> Gambar diset ke paragraf!")

def insert_image_before_para(doc, target_para, img_path, width_inches):
    """Sisipkan paragraf gambar baru tepat sebelum target_para."""
    new_para = doc.add_paragraph()
    new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = new_para.add_run()
    run.add_picture(img_path, width=Inches(width_inches))
    target_para._element.addprevious(new_para._element)
    print(f"  -> Gambar baru disisipkan sebelum caption!")

def find_caption_by_text(doc, text):
    """Cari paragraf Caption berdasarkan teks."""
    for i, p in enumerate(doc.paragraphs):
        if p.style.name == 'Caption' and text in p.text:
            return i, p
    return -1, None

# =========================================================
doc = Document(SRC)

print("=== STATUS GAMBAR DI SEKITAR CAPTION 4.1-4.3 ===")
for i in range(355, 370):
    p = doc.paragraphs[i]
    img_mark = ' [IMG]' if has_image(p) else ''
    print(f"[P{i}][{p.style.name}]{img_mark} | {p.text[:70]}")

print("\n=== OPERASI PERBAIKAN ===\n")

# ----------------------------------------------------------
# Gambar 4.1 Use Case Diagram
# -> Gambar ada di P359 (sebelum caption P360) - ganti isinya
# ----------------------------------------------------------
idx_41, cap_41 = find_caption_by_text(doc, "Gambar 4.1 Use Case Diagram")
if idx_41 >= 0:
    # Cari gambar sebelum caption
    img_para = None
    for j in range(idx_41-1, max(0, idx_41-5), -1):
        if has_image(doc.paragraphs[j]):
            img_para = doc.paragraphs[j]
            print(f"Gambar 4.1: ganti gambar di P{j}")
            clear_and_set_image(doc, img_para, IMG_USECASE, 5.5)
            break
    if img_para is None:
        print(f"Gambar 4.1: tidak ada gambar sebelum caption P{idx_41}, sisipkan...")
        insert_image_before_para(doc, cap_41, IMG_USECASE, 5.5)

# ----------------------------------------------------------
# Gambar 4.2 Activity Diagram Admin
# -> Tidak ada gambar sebelum caption P362 - sisipkan
# ----------------------------------------------------------
idx_42, cap_42 = find_caption_by_text(doc, "Gambar 4.2 Activity Diagram Admin")
if idx_42 >= 0:
    img_para = None
    for j in range(idx_42-1, max(0, idx_42-5), -1):
        if has_image(doc.paragraphs[j]):
            img_para = doc.paragraphs[j]
            break
        if doc.paragraphs[j].style.name == 'Caption':
            break  # Sudah lewat caption lain, stop
    if img_para is None:
        print(f"Gambar 4.2: sisipkan gambar baru sebelum caption P{idx_42}")
        insert_image_before_para(doc, cap_42, IMG_ACT_ADMIN, 5.5)
    else:
        print(f"Gambar 4.2: gambar sudah ada, ganti isinya")
        clear_and_set_image(doc, img_para, IMG_ACT_ADMIN, 5.5)

# ----------------------------------------------------------
# Gambar 4.3 Activity Diagram Kasir
# -> Gambar sudah ada di P364 (sebelum caption P365) - ganti isinya
# ----------------------------------------------------------
idx_43, cap_43 = find_caption_by_text(doc, "Gambar 4.3 Activity Diagram Kasir")
if idx_43 >= 0:
    img_para = None
    for j in range(idx_43-1, max(0, idx_43-5), -1):
        if has_image(doc.paragraphs[j]):
            img_para = doc.paragraphs[j]
            print(f"Gambar 4.3: ganti gambar di P{j}")
            clear_and_set_image(doc, img_para, IMG_ACT_KASIR, 5.5)
            break
        if doc.paragraphs[j].style.name == 'Caption':
            break
    if img_para is None:
        print(f"Gambar 4.3: tidak ada gambar sebelum caption P{idx_43}, sisipkan...")
        insert_image_before_para(doc, cap_43, IMG_ACT_KASIR, 5.5)

# ----------------------------------------------------------
# Gambar 4.19 Class Diagram
# -> Gambar ada di P418 (sebelum caption P419 setelah pindah) - ganti isinya
# ----------------------------------------------------------
idx_419, cap_419 = find_caption_by_text(doc, "Gambar 4.19 Class Diagram")
if idx_419 >= 0:
    img_para = None
    for j in range(idx_419-1, max(0, idx_419-5), -1):
        if has_image(doc.paragraphs[j]):
            img_para = doc.paragraphs[j]
            print(f"Gambar 4.19: ganti gambar di P{j}")
            clear_and_set_image(doc, img_para, IMG_CLASS, 5.5)
            break
    if img_para is None:
        print(f"Gambar 4.19: sisipkan gambar baru sebelum caption P{idx_419}")
        insert_image_before_para(doc, cap_419, IMG_CLASS, 5.5)

# ----------------------------------------------------------
# Gambar 4.20 ERD Database
# -> Gambar ada di P423 (sebelum caption P424) - ganti isinya
# ----------------------------------------------------------
idx_420, cap_420 = find_caption_by_text(doc, "Gambar 4.20 ERD Database")
if idx_420 >= 0:
    img_para = None
    for j in range(idx_420-1, max(0, idx_420-5), -1):
        if has_image(doc.paragraphs[j]):
            img_para = doc.paragraphs[j]
            print(f"Gambar 4.20: ganti gambar di P{j}")
            clear_and_set_image(doc, img_para, IMG_ERD, 5.5)
            break
    if img_para is None:
        print(f"Gambar 4.20: sisipkan gambar baru sebelum caption P{idx_420}")
        insert_image_before_para(doc, cap_420, IMG_ERD, 5.5)

# =========================================================
print(f"\nMenyimpan ke {DST}...")
doc.save(DST)
print("SELESAI!")

# VERIFIKASI FINAL
print("\n=== VERIFIKASI FINAL LENGKAP ===")
doc2 = Document(DST)
targets = ["Gambar 4.1", "Gambar 4.2", "Gambar 4.3", "Gambar 4.19", "Gambar 4.20"]
for t in targets:
    for i, p in enumerate(doc2.paragraphs):
        if p.style.name == 'Caption' and t in p.text:
            # Cek gambar sebelum caption
            has_img_before = False
            for j in range(i-1, max(0, i-5), -1):
                if has_image(doc2.paragraphs[j]):
                    has_img_before = True
                    print(f"  {t}: Gambar di P{j}, Caption di P{i} ✓")
                    break
                if doc2.paragraphs[j].style.name == 'Caption':
                    break
            if not has_img_before:
                print(f"  {t}: TIDAK ADA gambar sebelum caption P{i} ✗")
            break
