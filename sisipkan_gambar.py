
"""
Script menyisipkan gambar diagram baru setelah caption yang sesuai.
Menggunakan XML manipulation untuk menambahkan gambar secara tepat.
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from lxml import etree
import os
import copy

SRC = r'c:\Users\ASUS\Documents\Codex\Losarijaya2\TA2_14694_REVISI_v2.docx'
DST = r'c:\Users\ASUS\Documents\Codex\Losarijaya2\TA2_14694_REVISI_v2.docx'

ARTIFACT_DIR = r'C:\Users\ASUS\.gemini\antigravity-ide\brain\8d8f779d-6960-4d6b-8d92-000b7d2c5da9'
files = os.listdir(ARTIFACT_DIR)

IMG_USECASE   = os.path.join(ARTIFACT_DIR, [f for f in files if 'usecase_diagram' in f][0])
IMG_ACT_ADMIN = os.path.join(ARTIFACT_DIR, [f for f in files if 'activity_diagram_admin' in f][0])
IMG_ACT_KASIR = os.path.join(ARTIFACT_DIR, [f for f in files if 'activity_diagram_kasir' in f][0])
IMG_CLASS     = os.path.join(ARTIFACT_DIR, [f for f in files if 'class_diagram_sistem' in f][0])
IMG_ERD       = os.path.join(ARTIFACT_DIR, [f for f in files if 'erd_diagram_sistem' in f][0])

doc = Document(SRC)

def add_image_paragraph_after(doc, after_para_idx, img_path, width_inches=5.2):
    """
    Tambahkan paragraf gambar baru setelah paragraf di index tertentu.
    Gambar ditempatkan di tengah halaman.
    """
    target_para = doc.paragraphs[after_para_idx]
    
    # Buat paragraph baru dengan gambar
    new_para = doc.add_paragraph()
    new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = new_para.add_run()
    run.add_picture(img_path, width=Inches(width_inches))
    
    # Pindahkan paragraph baru setelah target menggunakan XML
    target_para._element.addnext(new_para._element)
    
    return new_para

def find_caption_para(doc, text_fragment):
    """Cari paragraf caption berdasarkan teks."""
    for i, p in enumerate(doc.paragraphs):
        if text_fragment in p.text:
            return i
    return -1

# =============================================================================
# SISIPKAN GAMBAR
# =============================================================================
print("=== MENYISIPKAN GAMBAR BARU ===\n")

# Gambar disisipkan setelah caption, sebelum deskripsi
# Format: (caption_text, image_path, width, label)
inserts = [
    # Setelah Gambar 4.1 caption (P355), sebelum penjelasan (P356)
    # Kita sisipkan SETELAH caption (P355)
    ("Gambar 4.1 Use Case Diagram", IMG_USECASE, 5.5, "Use Case Diagram"),
    # Setelah Gambar 4.2 caption (P358)
    ("Gambar 4.2 Activity Diagram Admin", IMG_ACT_ADMIN, 5.5, "Activity Admin"),
    # Setelah Gambar 4.3 caption (P361)
    ("Gambar 4.3 Activity Diagram Kasir", IMG_ACT_KASIR, 5.5, "Activity Kasir"),
    # Setelah Gambar 4.19 caption (P416)
    ("Gambar 4.19 Class Diagram", IMG_CLASS, 5.5, "Class Diagram"),
    # Setelah Gambar 4.20 caption (P421)
    ("Gambar 4.20 ERD Database", IMG_ERD, 5.5, "ERD Diagram"),
]

# Lakukan dalam urutan terbalik agar index tidak bergeser
inserts_with_idx = []
for caption_text, img_path, width, label in inserts:
    idx = find_caption_para(doc, caption_text)
    if idx >= 0:
        inserts_with_idx.append((idx, img_path, width, label, caption_text))
        print(f"  Ditemukan '{caption_text}' di P{idx}")
    else:
        print(f"  WARN: '{caption_text}' tidak ditemukan!")

# Urutkan dari index terbesar ke terkecil agar penambahan tidak menggeser posisi
inserts_with_idx.sort(key=lambda x: x[0], reverse=True)

for idx, img_path, width, label, caption in inserts_with_idx:
    print(f"\nMenyisipkan {label} setelah P{idx}...")
    add_image_paragraph_after(doc, idx, img_path, width)
    print(f"  -> OK!")

# =============================================================================
# SIMPAN
# =============================================================================
print(f"\nMenyimpan ke {DST}...")
doc.save(DST)
print("SELESAI!")

# Verifikasi
doc2 = Document(DST)
count_images = 0
for p in doc2.paragraphs:
    d1 = p._element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline')
    d2 = p._element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}anchor')
    if d1 or d2:
        count_images += 1

print(f"\nVerifikasi: Paragraf dengan gambar di dokumen hasil = {count_images}")
