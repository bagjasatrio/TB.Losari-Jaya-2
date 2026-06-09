
"""
Script yang BENAR untuk mengganti gambar di BAB 4.
Pola dokumen: [Gambar] -> [Caption] -> [Deskripsi]

Target penggantian:
- P359 (gambar sebelum caption 4.1) -> Use Case Diagram baru
- P362 (gambar antara caption 4.1 dan 4.2) -> Activity Admin baru  TAPI INI GAMBAR DUPLIKAT KITA!
- P365 (gambar antara caption 4.2 dan 4.3) -> Activity Kasir baru  TAPI INI GAMBAR DUPLIKAT KITA!
- P420 (gambar antara caption 4.18 dan 4.19) -> Class Diagram baru
- P425 (gambar antara caption 4.19 dan 4.20) -> ERD baru

Jadi: 
1. P359 = gambar asli Gambar 4.1 -> ganti dengan Use Case Diagram baru
2. P362 = gambar yang kita sisipkan salah posisi (activity admin) -> HAPUS, lalu
3. Gambar asli Gambar 4.2 ada di mana? Perlu cek kembali setelah hapus P362
4. Demikian seterusnya
"""

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from lxml import etree
import os
import copy

# Selalu mulai dari DOKUMEN ASLI untuk menghindari masalah duplikasi
SRC = r'c:\Users\ASUS\Documents\Codex\Losarijaya2\TA2_14694_REVISI.docx'
DST = r'c:\Users\ASUS\Documents\Codex\Losarijaya2\TA2_14694_REVISI_v2.docx'

ARTIFACT_DIR = r'C:\Users\ASUS\.gemini\antigravity-ide\brain\8d8f779d-6960-4d6b-8d92-000b7d2c5da9'
files = os.listdir(ARTIFACT_DIR)

IMG_USECASE   = os.path.join(ARTIFACT_DIR, [f for f in files if 'usecase_diagram' in f][0])
IMG_ACT_ADMIN = os.path.join(ARTIFACT_DIR, [f for f in files if 'activity_diagram_admin' in f][0])
IMG_ACT_KASIR = os.path.join(ARTIFACT_DIR, [f for f in files if 'activity_diagram_kasir' in f][0])
IMG_CLASS     = os.path.join(ARTIFACT_DIR, [f for f in files if 'class_diagram_sistem' in f][0])
IMG_ERD       = os.path.join(ARTIFACT_DIR, [f for f in files if 'erd_diagram_sistem' in f][0])

# ============================================================
# BUKA DOKUMEN YANG SUDAH DIREVISI (v2) YANG ADA TEKS REVISIANNYA
# TAPI KEMBALIKAN GAMBAR KE POSISI YANG BENAR
# ============================================================
print("Membuka v2 (dokumen dengan teks revisi)...")
doc = Document(DST)

# =============================================================
# STEP 1: Temukan paragraf gambar asli di BAB 4 (sebelum caption)
# dan hapus gambar duplikat yang salah posisi
# =============================================================

print("\n=== STEP 1: Analisis posisi gambar ===")
for i, p in enumerate(doc.paragraphs):
    d1 = p._element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline')
    d2 = p._element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}anchor')
    if (d1 or d2) and i > 330 and i < 450:
        next_text = ''
        if i < len(doc.paragraphs)-1:
            next_text = doc.paragraphs[i+1].text[:50]
        prev_text = ''
        if i > 0:
            prev_text = doc.paragraphs[i-1].text[:50]
        print(f"[P{i}] GAMBAR | prev: '{prev_text}' | next: '{next_text}'")

