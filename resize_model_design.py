
"""
Perbaikan layout BAB 4 Model Design:
- Kecilkan gambar halaman (Gambar 4.4 s/d 4.17) dari ~4.2in menjadi 3.0in
  agar caption + deskripsi ikut muat di halaman yang sama (2 set per halaman)
- Kurangi space_before/space_after pada paragraf gambar dan caption
- Gambar 4.1-4.3 (diagram UML) dan 4.18-4.20 dibiarkan lebih besar (sudah oke)
"""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from lxml import etree

SRC = r'c:\Users\ASUS\Documents\Codex\Losarijaya2\TA2_14694_REVISI_v2.docx'
DST = r'c:\Users\ASUS\Documents\Codex\Losarijaya2\TA2_14694_REVISI_v2.docx'

NS_WP_I  = '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline'
NS_A_EXT = '{http://schemas.openxmlformats.org/drawingml/2006/main}ext'
EMU_PER_INCH = 914400

def get_img_size_emu(p):
    """Return (cx, cy) dalam EMU, atau (None, None)."""
    for inline in p._element.findall(f'.//{NS_WP_I}'):
        ext = inline.find(f'.//{NS_A_EXT}')
        if ext is not None:
            return int(ext.get('cx', 0)), int(ext.get('cy', 0))
    return None, None

def resize_image_in_para(p, new_width_inches):
    """Ubah ukuran gambar secara proporsional ke lebar baru."""
    for inline in p._element.findall(f'.//{NS_WP_I}'):
        ext = inline.find(f'.//{NS_A_EXT}')
        if ext is not None:
            old_cx = int(ext.get('cx', 0))
            old_cy = int(ext.get('cy', 0))
            if old_cx == 0:
                continue
            new_cx = int(new_width_inches * EMU_PER_INCH)
            ratio = new_cx / old_cx
            new_cy = int(old_cy * ratio)
            ext.set('cx', str(new_cx))
            ext.set('cy', str(new_cy))
            # Juga update distT/distB di inline jika ada
            return old_cx / EMU_PER_INCH, old_cy / EMU_PER_INCH, new_cx / EMU_PER_INCH, new_cy / EMU_PER_INCH
    return None, None, None, None

def set_para_spacing(para, space_before_pt=0, space_after_pt=4):
    """Set space before/after paragraf dalam poin."""
    pf = para.paragraph_format
    pf.space_before = Pt(space_before_pt)
    pf.space_after = Pt(space_after_pt)

# ================================================================
doc = Document(SRC)
print("=== RESIZE GAMBAR MODEL DESIGN (Gambar 4.4 - 4.17) ===\n")

# Target: gambar halaman UI dikecilkan ke 3.3 inch lebar
# agar sepasang gambar bisa muat dalam 1-2 halaman dengan teksnya
NEW_UI_WIDTH = 3.3   # inch - untuk gambar halaman aplikasi (4.4-4.17)
NEW_SEQ_WIDTH = 5.0  # inch - untuk sequence diagram (4.18)

# Petakan caption ke ukuran target
# Format: teks yang ada di caption -> lebar baru dalam inch
caption_size_map = {
    "Gambar 4.4":  NEW_UI_WIDTH,
    "Gambar 4.5":  NEW_UI_WIDTH,
    "Gambar 4.6":  NEW_UI_WIDTH,
    "Gambar 4.7":  NEW_UI_WIDTH,
    "Gambar 4.8":  NEW_UI_WIDTH,
    "Gambar 4.9":  NEW_UI_WIDTH,
    "Gambar 4.10": NEW_UI_WIDTH,
    "Gambar 4.11": NEW_UI_WIDTH,
    "Gambar 4.12": NEW_UI_WIDTH,
    "Gambar 4.13": NEW_UI_WIDTH,
    "Gambar 4.14": NEW_UI_WIDTH,
    "Gambar 4.15": NEW_UI_WIDTH,
    "Gambar 4.16": NEW_UI_WIDTH,
    "Gambar 4.17": NEW_UI_WIDTH,
    "Gambar 4.18": NEW_SEQ_WIDTH,
}

# Cari semua paragraf gambar yang berhubungan dengan caption-caption di atas
# Pola: [Gambar] -> [Caption] -> [Deskripsi]
processed = set()

for i, p in enumerate(doc.paragraphs):
    if p.style.name != 'Caption':
        continue
    
    # Cek apakah caption ini masuk target
    target_width = None
    for cap_key, width in caption_size_map.items():
        if cap_key in p.text:
            target_width = width
            break
    if target_width is None:
        continue
    
    # Cari paragraf gambar SEBELUM caption (max 3 paragraf ke atas)
    img_para_idx = -1
    for j in range(i-1, max(0, i-4), -1):
        cx, cy = get_img_size_emu(doc.paragraphs[j])
        if cx is not None:
            img_para_idx = j
            break
        if doc.paragraphs[j].style.name == 'Caption':
            break
    
    if img_para_idx >= 0 and img_para_idx not in processed:
        img_para = doc.paragraphs[img_para_idx]
        old_w, old_h, new_w, new_h = resize_image_in_para(img_para, target_width)
        processed.add(img_para_idx)
        
        print(f"  {p.text[:40]:40s} | P{img_para_idx}: {old_w:.2f}in -> {new_w:.2f}in")
        
        # Kurangi spacing pada paragraf gambar (kurangi ruang kosong)
        set_para_spacing(img_para, space_before_pt=6, space_after_pt=2)
        
        # Kurangi spacing pada caption
        set_para_spacing(p, space_before_pt=0, space_after_pt=4)
        
        # Kurangi spacing pada paragraf deskripsi (setelah caption)
        if i+1 < len(doc.paragraphs):
            desc_para = doc.paragraphs[i+1]
            if desc_para.style.name in ['Normal', 'Normal (Web)']:
                set_para_spacing(desc_para, space_before_pt=0, space_after_pt=6)

print(f"\nTotal gambar diubah ukurannya: {len(processed)}")

# ================================================================
# TAMBAHAN: Set paragraph gambar agar rata tengah
# ================================================================
print("\n=== SET ALIGNMENT GAMBAR ===")
for i, p in enumerate(doc.paragraphs):
    cx, cy = get_img_size_emu(p)
    if cx and 365 < i < 415:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ================================================================
# SIMPAN
# ================================================================
print(f"\nMenyimpan ke {DST}...")
doc.save(DST)
print("SELESAI!")

# Verifikasi ukuran baru
print("\n=== VERIFIKASI UKURAN SETELAH RESIZE ===")
doc2 = Document(DST)
for i, p in enumerate(doc2.paragraphs):
    if 365 < i < 415:
        cx, cy = get_img_size_emu(p)
        if cx:
            cap_text = doc2.paragraphs[i+1].text[:40] if i+1 < len(doc2.paragraphs) else ''
            print(f"  [P{i}] W={cx/EMU_PER_INCH:.2f}in H={cy/EMU_PER_INCH:.2f}in -> {cap_text!r}")
