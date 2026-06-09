
"""
Script final: Ganti gambar diagram di posisi YANG BENAR di BAB 4.
Strategi:
  1. Buka dokumen teks-revisi (v2 yang ada teks bagus)
  2. Hapus paragraf gambar duplikat yang salah posisi (P362, P365 di v2)
  3. Ganti gambar di posisi ASLI yang benar:
     - P359 (sebelum caption 4.1) -> Use Case Diagram baru
     - P362-asli (sebelum caption 4.2 setelah hapus duplikat) -> Activity Admin
     - P365-asli (sebelum caption 4.3 setelah hapus duplikat) -> Activity Kasir
     - P420 (sebelum caption 4.19) -> Class Diagram baru
     - P425 (sebelum caption 4.20) -> ERD baru
"""

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from lxml import etree
import os
import io

SRC = r'c:\Users\ASUS\Documents\Codex\Losarijaya2\TA2_14694_REVISI_v2.docx'
DST = r'c:\Users\ASUS\Documents\Codex\Losarijaya2\TA2_14694_REVISI_v2.docx'

ARTIFACT_DIR = r'C:\Users\ASUS\.gemini\antigravity-ide\brain\8d8f779d-6960-4d6b-8d92-000b7d2c5da9'
files = os.listdir(ARTIFACT_DIR)

IMG_USECASE   = os.path.join(ARTIFACT_DIR, [f for f in files if 'usecase_diagram' in f][0])
IMG_ACT_ADMIN = os.path.join(ARTIFACT_DIR, [f for f in files if 'activity_diagram_admin' in f][0])
IMG_ACT_KASIR = os.path.join(ARTIFACT_DIR, [f for f in files if 'activity_diagram_kasir' in f][0])
IMG_CLASS     = os.path.join(ARTIFACT_DIR, [f for f in files if 'class_diagram_sistem' in f][0])
IMG_ERD       = os.path.join(ARTIFACT_DIR, [f for f in files if 'erd_diagram_sistem' in f][0])

# ===================================================================
# HELPER: Ganti isi gambar dalam paragraf yang sudah ada drawing
# ===================================================================

def replace_image_in_para(doc, para, new_img_path, width_inches=5.2):
    """
    Ganti gambar dalam paragraf yang sudah berisi drawing.
    Hapus semua drawing lama, tambahkan gambar baru.
    """
    p_elem = para._element
    
    # Hapus semua elemen drawing (inline dan anchor)
    NS_WD = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
    NS_WP = '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}'
    
    # Hapus semua run yang mengandung drawing
    runs_to_clear = []
    for run_elem in p_elem.findall(f'{NS_WD}r'):
        drawings = run_elem.findall(f'.//{NS_WP}inline') + run_elem.findall(f'.//{NS_WP}anchor')
        if drawings:
            runs_to_clear.append(run_elem)
    
    for run_elem in runs_to_clear:
        p_elem.remove(run_elem)
    
    # Tambah gambar baru
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.add_picture(new_img_path, width=Inches(width_inches))
    return True


def remove_paragraph(doc, para):
    """Hapus sebuah paragraf dari dokumen."""
    p_elem = para._element
    p_elem.getparent().remove(p_elem)


def find_para_with_image_before_caption(doc, caption_text, start_idx=0):
    """
    Cari paragraf yang berisi gambar yang posisinya tepat sebelum caption.
    Return (image_para_idx, caption_para_idx) atau (-1, -1)
    """
    for i in range(start_idx, len(doc.paragraphs)):
        p = doc.paragraphs[i]
        if p.style.name == 'Caption' and caption_text in p.text:
            # Cari gambar di paragraf sebelumnya
            for j in range(i-1, max(0, i-5), -1):
                prev = doc.paragraphs[j]
                d1 = prev._element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline')
                d2 = prev._element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}anchor')
                if d1 or d2:
                    return j, i
            return -1, i
    return -1, -1


def find_spurious_image_after_caption(doc, caption_idx):
    """
    Cari gambar duplikat yang ada setelah caption (posisi salah).
    Return idx atau -1
    """
    for i in range(caption_idx+1, min(caption_idx+5, len(doc.paragraphs))):
        p = doc.paragraphs[i]
        if p.style.name == 'Caption':
            break
        d1 = p._element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline')
        d2 = p._element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}anchor')
        if d1 or d2:
            return i
    return -1

# ===================================================================
# PROSES UTAMA
# ===================================================================

print("Membuka dokumen v2...")
doc = Document(SRC)

print("\n=== LANGKAH 1: Hapus gambar duplikat (yang salah posisi, setelah caption) ===")

# Gambar duplikat yang disisipkan di v1 sisipkan_gambar.py ada:
# - Setelah caption Gambar 4.1 (antara deskripsi 4.1 dan caption 4.2)
# - Setelah caption Gambar 4.2 (antara deskripsi 4.2 dan caption 4.3)
# - Setelah caption Gambar 4.19 (antara deskripsi 4.19 dan heading Implementasi Database)  
# - Setelah caption Gambar 4.20 (antara Tabel 4.1 dan implementasi)

# Kita cari dan hapus gambar-gambar di posisi salah ini
# (gambar yang muncul SETELAH deskripsi caption, bukan sebelum caption)

captions_with_spurious = [
    "Gambar 4.1 Use Case Diagram",
    "Gambar 4.2 Activity Diagram Admin",
]

for cap_text in captions_with_spurious:
    _, cap_idx = find_para_with_image_before_caption(doc, cap_text)
    if cap_idx >= 0:
        spurious_idx = find_spurious_image_after_caption(doc, cap_idx)
        if spurious_idx >= 0:
            print(f"  Menghapus gambar duplikat di P{spurious_idx} (setelah caption '{cap_text}' di P{cap_idx})")
            remove_paragraph(doc, doc.paragraphs[spurious_idx])
            print(f"  -> Dihapus!")
        else:
            print(f"  Tidak ada gambar duplikat setelah '{cap_text}'")
    else:
        print(f"  WARN: Caption '{cap_text}' tidak ditemukan!")

# Simpan sementara setelah hapus duplikat
print("\nMenyimpan setelah hapus duplikat...")
doc.save(DST)

# Buka ulang untuk refresh indeks
print("Membuka ulang dokumen...")
doc = Document(DST)

print("\n=== LANGKAH 2: Ganti gambar di posisi BENAR ===")

# Pemetaan: teks caption -> gambar baru
replacements = [
    ("Gambar 4.1 Use Case Diagram",      IMG_USECASE,   5.5, "Use Case Diagram"),
    ("Gambar 4.2 Activity Diagram Admin", IMG_ACT_ADMIN, 5.5, "Activity Admin"),
    ("Gambar 4.3 Activity Diagram Kasir", IMG_ACT_KASIR, 5.5, "Activity Kasir"),
    ("Gambar 4.19 Class Diagram",         IMG_CLASS,     5.5, "Class Diagram"),
    ("Gambar 4.20 ERD Database",          IMG_ERD,       5.5, "ERD Diagram"),
]

for cap_text, img_path, width, label in replacements:
    img_idx, cap_idx = find_para_with_image_before_caption(doc, cap_text)
    if cap_idx >= 0:
        if img_idx >= 0:
            print(f"\n  '{label}': gambar di P{img_idx}, caption di P{cap_idx}")
            para = doc.paragraphs[img_idx]
            replace_image_in_para(doc, para, img_path, width)
            print(f"  -> Gambar diganti dengan {label}!")
        else:
            print(f"\n  '{label}': caption di P{cap_idx}, TIDAK ADA gambar sebelumnya!")
            print(f"  -> Menyisipkan paragraf gambar baru sebelum caption...")
            # Sisipkan paragraf gambar sebelum caption
            cap_para = doc.paragraphs[cap_idx]
            new_para = doc.add_paragraph()
            new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = new_para.add_run()
            run.add_picture(img_path, width=Inches(width))
            # Pindahkan sebelum caption
            cap_para._element.addprevious(new_para._element)
            print(f"  -> Gambar {label} disisipkan sebelum caption!")
    else:
        print(f"\n  WARN: Caption '{cap_text}' tidak ditemukan!")

# ===================================================================
# SIMPAN FINAL
# ===================================================================
print(f"\n\nMenyimpan ke {DST}...")
doc.save(DST)
print("SELESAI!")

# Verifikasi akhir
print("\n=== VERIFIKASI AKHIR ===")
doc2 = Document(DST)
for i, p in enumerate(doc2.paragraphs):
    d1 = p._element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline')
    d2 = p._element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}anchor')
    if (d1 or d2) and 330 < i < 450:
        next_text = doc2.paragraphs[i+1].text[:50] if i < len(doc2.paragraphs)-1 else ''
        print(f"[P{i}] GAMBAR -> next: '{next_text}'")
