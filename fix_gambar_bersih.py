
"""
Script BERSIH: Mulai dari dokumen asli + teks revisi, ganti gambar dengan benar.
Pendekatan: Rebuild dari dokumen asli, terapkan revisi teks, lalu ganti gambar satu per satu.
"""

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# Gunakan dokumen dengan teks revisi yang sudah benar, tapi gambarnya masih asli
# Kita buat ulang dari dokumen teks revisi yang bersih
SRC_REVISI  = r'c:\Users\ASUS\Documents\Codex\Losarijaya2\TA2_14694_REVISI_v2.docx'
DST         = r'c:\Users\ASUS\Documents\Codex\Losarijaya2\TA2_14694_REVISI_v2.docx'

ARTIFACT_DIR = r'C:\Users\ASUS\.gemini\antigravity-ide\brain\8d8f779d-6960-4d6b-8d92-000b7d2c5da9'
files = os.listdir(ARTIFACT_DIR)

DIAGRAMS = {
    "Gambar 4.1 Use Case Diagram":      (os.path.join(ARTIFACT_DIR, [f for f in files if 'usecase_diagram' in f][0]),   5.5),
    "Gambar 4.2 Activity Diagram Admin": (os.path.join(ARTIFACT_DIR, [f for f in files if 'activity_diagram_admin' in f][0]), 5.5),
    "Gambar 4.3 Activity Diagram Kasir": (os.path.join(ARTIFACT_DIR, [f for f in files if 'activity_diagram_kasir' in f][0]), 5.5),
    "Gambar 4.19 Class Diagram":         (os.path.join(ARTIFACT_DIR, [f for f in files if 'class_diagram_sistem' in f][0]),  5.5),
    "Gambar 4.20 ERD Database":          (os.path.join(ARTIFACT_DIR, [f for f in files if 'erd_diagram_sistem' in f][0]),   5.5),
}

NS_WD = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
NS_WP_I = '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline'
NS_WP_A = '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}anchor'

def has_image(para):
    return bool(para._element.findall(f'.//{NS_WP_I}') or para._element.findall(f'.//{NS_WP_A}'))

def clear_image_from_para(para):
    """Hapus semua drawing dari paragraf."""
    p_elem = para._element
    for run_elem in list(p_elem.findall(f'{NS_WD}r')):
        if run_elem.findall(f'.//{NS_WP_I}') or run_elem.findall(f'.//{NS_WP_A}'):
            p_elem.remove(run_elem)

def set_image_in_para(doc, para, img_path, width_inches):
    """Set gambar baru dalam paragraf (mengganti yang lama)."""
    clear_image_from_para(para)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.add_picture(img_path, width=Inches(width_inches))

def remove_paragraph(para):
    para._element.getparent().remove(para._element)

# ============================================================
print("Membuka dokumen...")
doc = Document(SRC_REVISI)

# ============================================================
# LANGKAH 1: Bersihkan dokumen dari sisa gambar duplikat
# (gambar yang ada di antara deskripsi-gambar dan caption berikutnya)
# ============================================================
print("\n=== LANGKAH 1: Temukan dan hapus paragraf gambar yang salah posisi ===")

# Kumpulkan semua caption BAB 4 dan paragraf gambar di sekitarnya
paras = doc.paragraphs

# Cari semua caption style 'Caption' yang mengandung 'Gambar 4.'
caption_positions = {}
for i, p in enumerate(paras):
    if p.style.name == 'Caption' and 'Gambar 4.' in p.text:
        for key in DIAGRAMS:
            if key in p.text:
                caption_positions[key] = i
                break

print(f"Caption ditemukan: {caption_positions}")

# Untuk setiap caption, cek apakah ada gambar SETELAH caption (posisi salah)
# Gambar yang benar harusnya SEBELUM caption
to_remove = []
for cap_key, cap_idx in caption_positions.items():
    # Scan max 3 paragraf setelah caption, sebelum caption berikutnya
    for j in range(cap_idx+1, min(cap_idx+5, len(paras))):
        p = paras[j]
        if p.style.name == 'Caption':
            break
        if has_image(p):
            to_remove.append((j, cap_key))
            print(f"  Gambar salah posisi di P{j} (setelah caption '{cap_key}' di P{cap_idx})")

# Hapus dalam urutan terbalik agar index tidak bergeser
to_remove.sort(key=lambda x: x[0], reverse=True)
for idx, label in to_remove:
    print(f"  Menghapus P{idx} ({label})")
    remove_paragraph(doc.paragraphs[idx])

print(f"  Total dihapus: {len(to_remove)}")

# Simpan dan buka ulang
doc.save(DST)
doc = Document(DST)
paras = doc.paragraphs
print("  Dokumen disimpan dan dibuka ulang.")

# ============================================================
# LANGKAH 2: Update caption_positions setelah perubahan
# ============================================================
caption_positions = {}
for i, p in enumerate(paras):
    if p.style.name == 'Caption' and 'Gambar 4.' in p.text:
        for key in DIAGRAMS:
            if key in p.text:
                caption_positions[key] = i
                break

print(f"\nCaption positions setelah cleanup: {caption_positions}")

# ============================================================
# LANGKAH 3: Untuk setiap caption, cari gambar sebelumnya
# dan ganti dengan diagram baru. Jika tidak ada gambar, sisipkan.
# ============================================================
print("\n=== LANGKAH 2: Ganti/sisipkan gambar di posisi BENAR (sebelum caption) ===")

# Proses dalam urutan TERBALIK agar index tidak bergeser saat insert
sorted_targets = sorted(caption_positions.items(), key=lambda x: x[1], reverse=True)

for cap_key, cap_idx in sorted_targets:
    img_path, width = DIAGRAMS[cap_key]
    
    # Cari gambar di paragraf sebelum caption (max 3 paragraf ke atas)
    found_img_idx = -1
    for j in range(cap_idx-1, max(0, cap_idx-4), -1):
        if has_image(doc.paragraphs[j]):
            found_img_idx = j
            break
    
    if found_img_idx >= 0:
        print(f"\n  '{cap_key}': Ganti gambar di P{found_img_idx}")
        set_image_in_para(doc, doc.paragraphs[found_img_idx], img_path, width)
        print(f"  -> OK!")
    else:
        print(f"\n  '{cap_key}': Tidak ada gambar sebelum P{cap_idx}. Menyisipkan...")
        # Buat paragraf gambar baru dan sisipkan sebelum caption
        cap_para = doc.paragraphs[cap_idx]
        new_para = doc.add_paragraph()
        new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = new_para.add_run()
        run.add_picture(img_path, width=Inches(width))
        cap_para._element.addprevious(new_para._element)
        print(f"  -> Disisipkan!")

# ============================================================
# SIMPAN FINAL
# ============================================================
print(f"\n\nMenyimpan ke {DST}...")
doc.save(DST)
print("SELESAI!")

# ============================================================
# VERIFIKASI
# ============================================================
print("\n=== VERIFIKASI FINAL ===")
doc3 = Document(DST)
for i, p in enumerate(doc3.paragraphs):
    if has_image(p) and 330 < i < 440:
        next_t = doc3.paragraphs[i+1].text[:60] if i+1 < len(doc3.paragraphs) else ''
        prev_t = doc3.paragraphs[i-1].text[:60] if i > 0 else ''
        print(f"[P{i}] GAMBAR | next: '{next_t}'")
