
"""
Script untuk menyisipkan gambar baru ke dalam dokumen yang sudah direvisi.
Gambar disisipkan SETELAH caption yang sesuai, menggantikan atau melengkapi posisi gambar.
"""

import docx
from docx import Document
from docx.shared import Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from lxml import etree
import os
import copy

SRC = r'c:\Users\ASUS\Documents\Codex\Losarijaya2\TA2_14694_REVISI_v2.docx'
DST = r'c:\Users\ASUS\Documents\Codex\Losarijaya2\TA2_14694_REVISI_v2.docx'

ARTIFACT_DIR = r'C:\Users\ASUS\.gemini\antigravity-ide\brain\8d8f779d-6960-4d6b-8d92-000b7d2c5da9'

files = os.listdir(ARTIFACT_DIR)
IMG_USECASE   = os.path.join(ARTIFACT_DIR, [f for f in files if 'usecase_diagram' in f][0])
IMG_ACT_ADMIN = os.path.join(ARTIFACT_DIR, [f for f in files if 'activity_diagram_admin' in f][0])
IMG_ACT_KASIR = os.path.join(ARTIFACT_DIR, [f for f in files if 'activity_diagram_kasir' in f][0])
IMG_CLASS     = os.path.join(ARTIFACT_DIR, [f for f in files if 'class_diagram_sistem' in f][0])
IMG_ERD       = os.path.join(ARTIFACT_DIR, [f for f in files if 'erd_diagram_sistem' in f][0])

doc = Document(SRC)

# =============================================================================
# ANALISIS DOKUMEN: Temukan semua paragraf dengan gambar
# =============================================================================
print("=== ANALISIS PARAGRAF DENGAN GAMBAR ===")
WNS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
for i, para in enumerate(doc.paragraphs):
    # Cek apakah ada drawing/gambar
    drawings = para._element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline')
    drawings2 = para._element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}anchor')
    if drawings or drawings2:
        print(f"[P{i}] GAMBAR DITEMUKAN | Teks: '{para.text[:60]}'")

print("\n=== ANALISIS CAPTION BAB 4 ===")
for i, para in enumerate(doc.paragraphs):
    if 'Gambar 4.' in para.text or 'Gambar 2.' in para.text:
        print(f"[P{i}][{para.style.name}] {para.text}")
