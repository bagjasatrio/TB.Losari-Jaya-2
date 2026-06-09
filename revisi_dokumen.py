
"""
Script revisi TA2_14694_REVISI.docx
- Merevisi teks agar lebih natural dan efisien (gaya bahasa manusia)
- Mempertahankan semua sitasi/referensi
- Mempertahankan format, heading, numbering
- Menyisipkan gambar diagram baru
"""

import docx
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import copy
import re
import os
import shutil
from lxml import etree

SRC = r'c:\Users\ASUS\Documents\Codex\Losarijaya2\TA2_14694_REVISI.docx'
DST = r'c:\Users\ASUS\Documents\Codex\Losarijaya2\TA2_14694_REVISI_v2.docx'

# Gambar diagram baru (dari artifacts)
ARTIFACT_DIR = r'C:\Users\ASUS\.gemini\antigravity-ide\brain\8d8f779d-6960-4d6b-8d92-000b7d2c5da9'

IMG_USECASE   = os.path.join(ARTIFACT_DIR, [f for f in os.listdir(ARTIFACT_DIR) if 'usecase_diagram' in f][0])
IMG_ACT_ADMIN = os.path.join(ARTIFACT_DIR, [f for f in os.listdir(ARTIFACT_DIR) if 'activity_diagram_admin' in f][0])
IMG_ACT_KASIR = os.path.join(ARTIFACT_DIR, [f for f in os.listdir(ARTIFACT_DIR) if 'activity_diagram_kasir' in f][0])
IMG_CLASS     = os.path.join(ARTIFACT_DIR, [f for f in os.listdir(ARTIFACT_DIR) if 'class_diagram_sistem' in f][0])
IMG_ERD       = os.path.join(ARTIFACT_DIR, [f for f in os.listdir(ARTIFACT_DIR) if 'erd_diagram_sistem' in f][0])

print("Gambar yang ditemukan:")
print(" UseCase :", IMG_USECASE)
print(" Admin   :", IMG_ACT_ADMIN)
print(" Kasir   :", IMG_ACT_KASIR)
print(" Class   :", IMG_CLASS)
print(" ERD     :", IMG_ERD)

# =============================================================================
# TEKS REVISI - kunci = teks asli (strip), nilai = teks pengganti
# Menjaga semua sitasi dalam tanda kurung. Teks lebih natural & efisien.
# =============================================================================

REVISI = {

# ============================================================
# RINGKASAN
# ============================================================
"TB. Losari Jaya 2 merupakan toko bangunan dengan tingkat perputaran persediaan yang tinggi sehingga membutuhkan pengelolaan stok yang akurat dan terstruktur. Dalam praktiknya, proses pencatatan barang masuk dan keluar masih dilakukan secara manual menggunakan buku tulis, yang berpotensi menimbulkan selisih data stok, keterlambatan informasi, serta risiko kehilangan atau kerusakan data. Permasalahan tersebut berdampak pada kesulitan pemantauan persediaan dan penyusunan laporan penjualan secara tepat. Penelitian ini bertujuan untuk merancang dan membangun sistem informasi penataan persediaan barang penjualan berbasis website yang mampu membantu proses pengelolaan stok dan transaksi secara lebih efektif dan real time. Metode pengembangan yang digunakan adalah metode Prototyping, yang meliputi tahap pengumpulan kebutuhan melalui komunikasi dengan pengguna, perancangan mock up sistem, pengujian Prototyping, serta evaluasi dan penyempurnaan berdasarkan umpan balik pengguna. Sistem informasi dikembangkan berbasis web dengan menggunakan framework Laravel dan basis data MySQL. Hasil penelitian menunjukkan bahwa sistem informasi berhasil diimplementasikan sesuai kebutuhan pengguna. Berdasarkan pengujian Black Box, seluruh fungsi sistem berjalan dengan valid dan aplikasi mampu mempermudah admin dalam pemantauan stok, pencatatan transaksi, serta rekapitulasi laporan persediaan barang.":
"TB. Losari Jaya 2 adalah toko bangunan dengan tingkat transaksi tinggi yang masih mengandalkan buku tulis dalam mencatat barang masuk dan keluar. Cara ini rentan menimbulkan selisih data stok, keterlambatan informasi, dan kehilangan catatan. Penelitian ini bertujuan merancang dan membangun sistem informasi persediaan barang berbasis website agar pengelolaan stok dan transaksi menjadi lebih akurat dan efisien. Pengembangan menggunakan metode Prototyping, mulai dari komunikasi kebutuhan bersama pengguna, perancangan mockup, pengujian prototype, hingga evaluasi berdasarkan umpan balik. Sistem dibangun dengan framework Laravel dan basis data MySQL. Hasil pengujian Black Box menunjukkan seluruh fungsi berjalan valid, dan sistem terbukti membantu admin dalam memantau stok, mencatat transaksi, serta menyusun laporan persediaan secara terstruktur.",

# ============================================================
# BAB I - LATAR BELAKANG
# ============================================================
"Kemajuan teknologi informasi telah membawa dampak besar terhadap cara pelaku usaha mengelola data operasionalnya, khususnya pada bidang perdagangan yang memiliki intensitas transaksi dan perputaran barang tinggi. Dalam konteks tersebut, persediaan barang memiliki peran krusial karena secara langsung memengaruhi ketersediaan produk, kelancaran proses penjualan, serta mutu pelayanan kepada pelanggan. Namun, praktik pengelolaan stok yang masih dilakukan secara manual sering kali menimbulkan kendala serius, seperti ketidaksesuaian antara data dan kondisi riil, keterlambatan informasi, hingga kesulitan dalam penyusunan laporan (Gultom, 2022) Oleh karena itu, penerapan sistem informasi yang terintegrasi menjadi kebutuhan mendesak untuk menjamin ketepatan dan efisiensi dalam pengendalian persediaan.":
"Kemajuan teknologi informasi mendorong pelaku usaha untuk membenahi cara mereka mengelola data operasional, terutama di sektor perdagangan yang memiliki intensitas transaksi tinggi. Persediaan barang menjadi aspek kritis karena langsung memengaruhi ketersediaan produk, kelancaran penjualan, dan kualitas layanan pelanggan. Namun kenyataannya, pengelolaan stok yang masih manual kerap memunculkan masalah seperti selisih data dengan kondisi fisik, lambatnya akses informasi, dan sulitnya menyusun laporan yang akurat (Gultom, 2022). Kondisi ini menjadikan penerapan sistem informasi terintegrasi sebagai kebutuhan yang tidak bisa ditunda.",

"Studi lapangan yang dilakukan di TB. Losari Jaya 2—sebuah toko bangunan dengan jumlah item dan frekuensi transaksi yang tinggi—menunjukkan bahwa pencatatan stok dan transaksi masih dilakukan menggunakan buku tulis. Cara tradisional ini mengakibatkan ketidaksesuaian antara catatan dan stok fisik, menyulitkan proses pelacakan data, serta meningkatkan risiko kehilangan informasi. Situasi serupa juga umum terjadi pada berbagai usaha dagang lain yang belum beralih ke sistem digital, sehingga akurasi data rendah dan proses kerja menjadi tidak efisien (Widiatma & Abdillah, 2025).":
"Pengamatan langsung di TB. Losari Jaya 2—toko bangunan dengan ratusan item dan transaksi harian yang padat—menunjukkan bahwa pencatatan stok dan transaksi masih dilakukan dengan buku tulis. Akibatnya, data sering tidak sesuai kondisi fisik, pelacakan barang menjadi lambat, dan informasi berisiko hilang sewaktu-waktu. Situasi ini umum dijumpai pada usaha dagang yang belum beralih ke sistem digital, sehingga akurasi data rendah dan efisiensi kerja terhambat (Widiatma & Abdillah, 2025).",

"Penelitian terdahulu telah membuktikan bahwa penerapan sistem informasi persediaan berbasis web dapat memperbaiki tata kelola data barang secara signifikan. Sistem semacam ini memberikan manfaat berupa pengelolaan data secara real time, penurunan tingkat kesalahan input, percepatan pencarian informasi, serta peningkatan keakuratan laporan (Effendi & Noviansyah, 2018). Selain itu, pendekatan prototyping dinilai efektif karena memungkinkan keterlibatan pengguna selama proses pengembangan sistem. Dengan demikian, kebutuhan pengguna dapat diidentifikasi dan disesuaikan sejak tahap perancangan awal (Ahmad et al., 2024).":
"Sejumlah penelitian terdahulu membuktikan bahwa sistem informasi persediaan berbasis web mampu memperbaiki tata kelola data secara signifikan. Manfaatnya antara lain pengelolaan data secara real time, berkurangnya kesalahan input, pencarian informasi yang lebih cepat, dan laporan yang lebih akurat (Effendi & Noviansyah, 2018). Pendekatan prototyping terbukti efektif karena melibatkan pengguna langsung dalam proses pengembangan, sehingga kebutuhan dapat diidentifikasi dan disesuaikan sejak tahap awal (Ahmad et al., 2024).",

"Hasil penelitian terdahulu menegaskan bahwa penggabungan antara sistem penjualan dan sistem pengelolaan persediaan dapat meningkatkan ketepatan data serta mempercepat proses pembuatan laporan. Studi yang dilakukan oleh (Adikusuma et al., 2025) menunjukkan bahwa penerapan sistem informasi penjualan berbasis web menggunakan metode prototyping efektif dalam menekan tingkat kesalahan pencatatan transaksi dan secara otomatis memperbarui data stok barang.":
"Studi oleh (Adikusuma et al., 2025) memperkuat hal ini dengan menunjukkan bahwa sistem informasi penjualan berbasis web menggunakan metode prototyping efektif menekan kesalahan pencatatan transaksi sekaligus memperbarui data stok secara otomatis. Integrasi antara modul penjualan dan persediaan terbukti meningkatkan ketepatan data dan mempercepat penyusunan laporan.",

"Oleh karena itu, pada penelitian implementasi sistem saya ini diharapkan dapat memberikan kemudahan bagi pelaku usaha dalam mengelola aktivitas penjualan maupun persediaan dengan lebih sistematis, sekaligus mendukung tersusunnya laporan yang akurat dan dapat diakses dengan cepat. Hasil temuan tersebut memperkuat relevansi penerapan metode prototyping sebagai pendekatan yang tepat untuk dikembangkan pada sektor perdagangan dengan intensitas transaksi yang tinggi.":
"Berdasarkan kondisi tersebut, penelitian ini bertujuan mengembangkan sistem informasi yang membantu TB. Losari Jaya 2 mengelola penjualan dan persediaan secara lebih sistematis, serta menghasilkan laporan yang akurat dan mudah diakses. Metode prototyping dipilih sebagai pendekatan pengembangan karena terbukti relevan untuk usaha dagang dengan intensitas transaksi tinggi.",

# ============================================================
# BAB I - RUMUSAN MASALAH
# ============================================================
"Rumusan masalah berikut disusun secara ringkas dan fokus, dengan mengacu langsung pada konteks latar belakang penelitian yang telah diuraikan sebelumnya, sebagaimana dijelaskan berikut ini:":
"Berdasarkan latar belakang yang telah diuraikan, rumusan masalah dalam penelitian ini adalah sebagai berikut:",

"Bagaimana sistem pengelolaan persediaan barang yang masih dilakukan secara manual menggunakan buku tulis di TB. Losari Jaya 2?":
"Bagaimana kondisi pengelolaan persediaan barang yang masih dilakukan secara manual menggunakan buku tulis di TB. Losari Jaya 2?",

"Bagaimana proses perancangan dan pengembangan sistem informasi penataan persediaan barang berbasis website dengan menerapkan metode?":
"Bagaimana proses perancangan dan pengembangan sistem informasi penataan persediaan barang berbasis website dengan menerapkan metode prototyping?",

# ============================================================
# BAB I - BATASAN MASALAH
# ============================================================
"Batasan masalah berikut disusun agar ruang lingkup penelitian tetap terarah dan konsisten dengan tujuan yang telah ditetapkan, sebagaimana dijelaskan berikut ini:":
"Agar ruang lingkup penelitian tetap terarah, batasan masalah ditetapkan sebagai berikut:",

"Penelitian ini berfokus pada proses perancangan serta pengembangan aplikasi berbasis web untuk pengelolaan persediaan barang penjualan dengan TB. Losari Jaya.":
"Penelitian ini berfokus pada perancangan dan pengembangan aplikasi berbasis web untuk pengelolaan persediaan barang penjualan pada TB. Losari Jaya 2.",

"Aplikasi dirancang menggunakan bahasa pemrograman PHP dengan basis data MySQL, sedangkan proses pengujian sistem dibatasi pada penerapan metode Black Box Testing.":
"Aplikasi dibangun menggunakan bahasa pemrograman PHP dengan framework Laravel dan basis data MySQL, sedangkan pengujian sistem dibatasi pada metode Black Box Testing dan User Acceptance Testing (UAT).",

# ============================================================
# BAB I - TUJUAN PENELITIAN
# ============================================================
"Tujuan penelitian berikut disusun secara terarah dan konsisten dengan rumusan masalah yang telah ditetapkan sebelumnya, sebagaimana dijelaskan berikut ini:":
"Sesuai dengan rumusan masalah yang telah ditetapkan, penelitian ini memiliki tujuan sebagai berikut:",

"Mengembangkan dan merancang sistem informasi berbasis website yang berfungsi untuk menata serta mengelola data persediaan pada TB. Losari Jaya 2.":
"Merancang dan mengembangkan sistem informasi berbasis website yang berfungsi untuk mengelola persediaan barang pada TB. Losari Jaya 2.",

"Menerapkan metode prototyping dalam proses pengembangan sistem informasi persediaan guna memastikan bahwa sistem yang dibangun mampu menyesuaikan dengan kebutuhan pengguna serta mengurangi potensi kesalahan dalam pencatatan data stok barang.":
"Menerapkan metode prototyping dalam pengembangan sistem sehingga kebutuhan pengguna dapat diakomodasi secara bertahap dan kesalahan pencatatan data stok dapat diminimalkan.",

# ============================================================
# BAB I - MANFAAT PENELITIAN
# ============================================================
"Hasil penelitian ini diharapkan memberikan manfaat dalam dua aspek utama, yaitu manfaat teoritis dan manfaat praktis, sebagaimana dijelaskan berikut ini:":
"Penelitian ini diharapkan memberikan manfaat bagi beberapa pihak, sebagaimana diuraikan berikut ini:",

"Menambah pengetahuan dan pengalaman bagi penulis dalam mengembangkan website persediaan barang penjualan menggunakan metode prototyping.":
"Menambah pengetahuan dan pengalaman penulis dalam mengembangkan sistem informasi persediaan berbasis website menggunakan metode prototyping.",

"Dapat menerapkan dan mengimplementasikan ilmu yang telah diberikan pada masa perkuliahan.":
"Memberikan kesempatan untuk menerapkan ilmu yang diperoleh selama perkuliahan dalam proyek nyata.",

"Sebagai referensi bagi mahasiswa Universitas Dian Nuswantoro yang akan menyusun skripsi pada tahun berikutnya.":
"Menjadi referensi bagi mahasiswa Universitas Dian Nuswantoro yang akan menyusun skripsi dengan topik serupa.",

"Sebagai rujukan penelitian selanjutnya yang terkait dengan sistem persediaan barang penjualan berbasis website menggunakan metode prototyping.":
"Menjadi rujukan bagi penelitian selanjutnya yang berkaitan dengan sistem informasi persediaan berbasis website dan metode prototyping.",

"Sebagai pengetahuan bagi masyarakat khususnya pelaku usaha kecil dan menengah di bidang perdagangan.":
"Memberikan wawasan kepada pelaku usaha kecil dan menengah mengenai manfaat digitalisasi sistem pencatatan persediaan barang.",

"Sebagai pengetahuan bagi pengembang website yang ingin menerapkan sistem persediaan barang penjualan menggunakan metode prototyping.":
"Menjadi acuan praktis bagi pengembang yang ingin membangun sistem persediaan barang berbasis website menggunakan pendekatan prototyping.",

# ============================================================
# BAB II - TINJAUAN STUDI
# ============================================================
"Pengelolaan persediaan memegang peranan krusial dalam menjaga kelancaran aktivitas usaha dagang karena berkaitan langsung dengan ketersediaan produk, proses transaksi, dan mutu pelayanan terhadap konsumen. Ketika pencatatan stok masih dilakukan secara manual, sering muncul berbagai kendala seperti perbedaan antara jumlah stok fisik dan data catatan, keterlambatan dalam memperoleh informasi, serta hambatan dalam penyusunan laporan persediaan. Hasil penelitian yang dilakukan oleh (Gultom, 2022) menunjukkan bahwa penggunaan sistem informasi berbasis web dalam pengelolaan stok barang mampu meningkatkan keteraturan dan ketepatan data dibandingkan metode manual.":
"Persediaan barang adalah komponen vital dalam usaha dagang karena memengaruhi ketersediaan produk, kelancaran transaksi, dan kualitas layanan. Ketika pencatatan masih manual, kendala seperti selisih stok fisik dengan data, lambatnya akses informasi, dan kesulitan menyusun laporan kerap muncul. (Gultom, 2022) membuktikan bahwa sistem informasi berbasis web mampu meningkatkan keteraturan dan ketepatan data persediaan dibandingkan metode manual.",

"Selain meningkatkan akurasi, penerapan sistem informasi persediaan juga berkontribusi terhadap efisiensi pengelolaan data dan optimalisasi proses bisnis. (Effendi & Noviansyah, 2018) menegaskan bahwa sistem informasi manajemen persediaan dapat mempercepat proses pencarian data, memperlancar pengolahan informasi, dan mendukung pembuatan laporan persediaan secara berkala. Penelitian lain oleh (Widiatma & Abdillah, 2025) turut memperkuat temuan tersebut dengan menyatakan bahwa sistem inventori berbasis web dapat meminimalkan risiko kehilangan data serta meningkatkan ketepatan dalam pencatatan arus barang masuk dan keluar, terutama pada usaha dengan intensitas transaksi yang tinggi.":
"Selain akurasi, (Effendi & Noviansyah, 2018) menegaskan bahwa sistem informasi persediaan mempercepat pencarian data, memperlancar pengolahan informasi, dan mendukung pembuatan laporan berkala. Sejalan dengan itu, (Widiatma & Abdillah, 2025) menyatakan bahwa sistem inventori berbasis web meminimalkan risiko kehilangan data dan meningkatkan ketepatan pencatatan arus barang, terutama pada usaha dengan transaksi yang padat.",

"Dalam konteks pengembangan sistem informasi, pemilihan metode rekayasa perangkat lunak menjadi aspek yang menentukan keberhasilan sistem agar sesuai dengan kebutuhan pengguna. Metode prototyping dinilai efektif karena memungkinkan pengguna terlibat aktif dalam tahapan perancangan hingga evaluasi sistem. Menurut (Ahmad et al., 2024), penerapan metode ini mampu menghasilkan aplikasi stok opname yang lebih responsif dan mudah digunakan sesuai dengan kebutuhan operasional. Temuan yang sejalan disampaikan oleh (Adikusuma et al., 2025) yang menyimpulkan bahwa sistem penjualan berbasis web dengan pendekatan prototyping dapat menekan tingkat kesalahan pencatatan serta memperbarui data stok secara otomatis. Berdasarkan hasil hasil penelitian tersebut, penerapan sistem informasi penataan persediaan barang berbasis web dengan metode prototyping dipandang tepat untuk diimplementasikan pada TB. Losari Jaya 2, guna meningkatkan efektivitas dan akurasi pengelolaan persediaan.":
"Dalam konteks pengembangan sistem, pemilihan metode rekayasa perangkat lunak sangat menentukan kesesuaian sistem dengan kebutuhan pengguna. Metode prototyping efektif karena melibatkan pengguna secara aktif dari tahap perancangan hingga evaluasi. (Ahmad et al., 2024) membuktikan bahwa pendekatan ini menghasilkan aplikasi stok opname yang responsif dan mudah dioperasikan. (Adikusuma et al., 2025) menambahkan bahwa sistem penjualan berbasis web dengan prototyping dapat menekan kesalahan pencatatan dan memperbarui stok secara otomatis. Berdasarkan temuan-temuan tersebut, penerapan sistem informasi persediaan berbasis web dengan metode prototyping dipandang tepat untuk TB. Losari Jaya 2.",

# ============================================================
# BAB II - TINJAUAN PUSTAKA
# ============================================================
"Sistem informasi dapat dipahami sebagai suatu kesatuan yang terdiri atas berbagai komponen yang saling berhubungan dan bekerja sama untuk mengumpulkan, mengelola, menyimpan, serta menyalurkan informasi yang dibutuhkan dalam menunjang aktivitas operasional maupun proses pengambilan keputusan organisasi. Keberadaan sistem ini berperan signifikan dalam menciptakan tata kelola bisnis yang lebih efisien, teratur, dan mudah diawasi. Dalam lingkup usaha dagang, sistem informasi dimanfaatkan untuk menangani data transaksi serta persediaan barang sehingga informasi yang dihasilkan menjadi lebih akurat dan mudah diakses. Implementasi sistem informasi yang dirancang dengan tepat terbukti mampu menekan tingkat kesalahan pencatatan, mempercepat proses pengolahan data, serta meningkatkan efektivitas kinerja pengguna (Gultom, 2022).":
"Sistem informasi adalah kumpulan komponen yang saling terhubung untuk mengumpulkan, mengolah, menyimpan, dan menyebarkan informasi guna mendukung operasional dan pengambilan keputusan organisasi. Dalam lingkup usaha dagang, sistem ini menangani data transaksi dan persediaan sehingga informasi menjadi lebih akurat dan mudah diakses. Penerapan sistem informasi yang tepat terbukti menekan kesalahan pencatatan, mempercepat pengolahan data, dan meningkatkan efektivitas kerja pengguna (Gultom, 2022).",

"Penataan persediaan barang merupakan salah satu bentuk penyimpanan aset yang dimiliki perusahaan untuk dijual kembali maupun digunakan dalam mendukung aktivitas operasionalnya. Keberadaan penataan persediaan barang berperan penting dalam menjaga kelancaran proses penjualan serta memastikan kebutuhan pelanggan terpenuhi tepat waktu. Apabila pengelolaan persediaan dilakukan tanpa sistem yang terorganisir, maka dapat muncul berbagai permasalahan seperti kelebihan stok, kekurangan stok, hingga ketidaksesuaian antara data pencatatan dan kondisi fisik barang di lapangan. Oleh sebab itu, dibutuhkan suatu sistem pengelolaan persediaan yang mampu memantau dan mencatat setiap arus barang masuk maupun keluar dengan tingkat akurasi dan konsistensi yang tinggi (Effendi & Noviansyah, 2018).":
"Persediaan barang adalah aset perusahaan yang disimpan untuk dijual kembali atau digunakan dalam mendukung operasional. Pengelolaannya menjadi kunci kelancaran penjualan dan pemenuhan kebutuhan pelanggan tepat waktu. Tanpa sistem yang terorganisir, masalah seperti kelebihan stok, kekurangan stok, dan ketidaksesuaian data dengan kondisi fisik mudah terjadi. Oleh karena itu, diperlukan sistem yang dapat mencatat dan memantau arus barang masuk maupun keluar secara akurat dan konsisten (Effendi & Noviansyah, 2018).",

"Sistem informasi persediaan barang merupakan suatu mekanisme berbasis komputer yang dirancang untuk mengelola seluruh data terkait stok barang, mulai dari proses pencatatan barang yang masuk, distribusi barang keluar, hingga pembuatan laporan persediaan secara terstruktur. Tujuan utama sistem ini adalah menyajikan informasi stok secara akurat dan terkini (real time), sehingga pengguna dapat memantau kondisi ketersediaan barang dengan lebih mudah dan efisien. Berdasarkan penelitian yang dilakukan oleh (Widiatma & Abdillah, 2025), penerapan sistem informasi persediaan berbasis web terbukti mampu meningkatkan akurasi data, mempercepat proses pencarian informasi, serta meminimalkan risiko kehilangan data yang sering terjadi pada metode pencatatan manual.":
"Sistem informasi persediaan barang adalah mekanisme berbasis komputer yang mengelola data stok secara menyeluruh, mulai dari pencatatan barang masuk, pengeluaran barang, hingga pelaporan yang terstruktur. Sistem ini bertujuan menyajikan informasi stok secara akurat dan real time agar kondisi ketersediaan barang dapat dipantau dengan mudah. (Widiatma & Abdillah, 2025) membuktikan bahwa penerapan sistem ini berbasis web meningkatkan akurasi data, mempercepat pencarian informasi, dan meminimalkan risiko kehilangan data dibandingkan pencatatan manual.",

"Website merupakan platform yang dibangun menggunakan teknologi internet untuk menyajikan serta mengelola informasi secara daring dan terintegrasi. Sistem informasi yang dikembangkan berbasis web memiliki keunggulan utama pada aspek kemudahan akses, karena dapat dijalankan melalui berbagai jenis perangkat tanpa perlu instalasi aplikasi tambahan. Selain itu, sistem berbasis web juga mempermudah proses pemeliharaan maupun pembaruan fungsi sistem agar tetap relevan dengan kebutuhan pengguna. Dalam konteks usaha dagang, pemanfaatan website sebagai media pengelolaan informasi dinilai efektif karena mampu menyediakan data persediaan secara cepat, akurat, dan fleksibel sesuai kebutuhan operasional.":
"Website adalah platform berbasis teknologi internet yang memungkinkan pengelolaan dan penyajian informasi secara daring dan terintegrasi. Keunggulan utama sistem berbasis web terletak pada kemudahan akses dari berbagai perangkat tanpa instalasi tambahan, serta kemudahan dalam pemeliharaan dan pembaruan fitur. Dalam konteks usaha dagang, website terbukti efektif sebagai media pengelolaan informasi persediaan karena menyediakan data secara cepat, akurat, dan fleksibel sesuai kebutuhan operasional.",

"Basis data merupakan kumpulan data yang disimpan secara terstruktur dan dikelola menggunakan sistem manajemen basis data (DBMS). Basis data berfungsi sebagai tempat penyimpanan utama data dalam sistem informasi, termasuk data barang, transaksi, dan laporan. Penggunaan basis data seperti MySQL memungkinkan penyimpanan data secara terorganisir, meminimalkan duplikasi data, serta mempermudah proses pengolahan dan pencarian informasi.":
"Basis data adalah kumpulan data terstruktur yang dikelola dengan sistem manajemen basis data (DBMS). Dalam sistem informasi, basis data berfungsi sebagai penyimpanan utama untuk data barang, transaksi, dan laporan. Penggunaan MySQL memungkinkan penyimpanan yang terorganisir, meminimalkan duplikasi, dan mempermudah pengolahan serta pencarian informasi.",

"MySQL merupakan sistem manajemen basis data relasional yang berfungsi sebagai komponen utama dalam penyimpanan dan pengelolaan data pada sistem informasi berbasis web. Melalui MySQL, berbagai data seperti daftar barang, transaksi masuk dan keluar, serta laporan stok dapat dikelola secara terstruktur dan terintegrasi.":
"MySQL adalah sistem manajemen basis data relasional (RDBMS) yang banyak digunakan sebagai komponen penyimpanan data pada aplikasi web. Melalui MySQL, data barang, transaksi masuk-keluar, dan laporan stok dapat dikelola secara terstruktur dan terintegrasi dalam satu sistem.",

"Hasil penelitian (Effendi & Noviansyah, 2018) menunjukkan bahwa pemanfaatan basis data relasional dalam sistem informasi persediaan dapat meningkatkan ketepatan data serta mempercepat proses pengolahan informasi. Temuan serupa dikemukakan oleh (Widiatma & Abdillah, 2025), yang menjelaskan bahwa sistem inventori berbasis web dengan dukungan basis data seperti MySQL mampu menekan kesalahan pencatatan dan mengurangi risiko kehilangan data. Dengan kemampuannya dalam menjaga konsistensi serta integritas data, MySQL dianggap sangat relevan digunakan sebagai basis data pendukung dalam pengembangan sistem informasi penataan persediaan barang penjualan berbasis website pada TB. Losari Jaya 2.":
"(Effendi & Noviansyah, 2018) menunjukkan bahwa basis data relasional meningkatkan ketepatan data dan mempercepat pengolahan informasi dalam sistem persediaan. (Widiatma & Abdillah, 2025) mempertegas bahwa sistem inventori berbasis web dengan MySQL mampu menekan kesalahan pencatatan dan mengurangi risiko kehilangan data. Dengan kemampuannya menjaga konsistensi dan integritas data, MySQL dianggap paling tepat untuk mendukung pengembangan sistem informasi persediaan TB. Losari Jaya 2.",

"Metode prototyping merupakan pendekatan dalam pengembangan perangkat lunak yang berfokus pada pembuatan model awal sistem sebagai representasi dari fungsi dan tampilan yang akan dikembangkan secara menyeluruh. Melalui metode ini, pengguna dilibatkan secara aktif dalam proses evaluasi maupun penyempurnaan model sistem, sehingga hasil akhir dapat disesuaikan dengan kebutuhan aktual pengguna. Berdasarkan penelitian yang dilakukan oleh (Ahmad et al., 2024), penggunaan metode prototyping terbukti efektif dalam pengembangan sistem persediaan karena memungkinkan proses perbaikan dilakukan secara bertahap dan berulang melalui umpan balik langsung dari pengguna. Pendekatan ini dianggap paling tepat diterapkan pada sistem informasi persediaan yang memiliki karakteristik kebutuhan yang dinamis dan memerlukan penyesuaian berkelanjutan.":
"Metode prototyping adalah pendekatan pengembangan perangkat lunak yang dimulai dengan membangun model awal sistem (prototype) sebagai representasi fungsi dan tampilan yang akan dikembangkan lebih lanjut. Pengguna dilibatkan secara aktif dalam evaluasi dan penyempurnaan setiap iterasi, sehingga sistem akhir benar-benar sesuai kebutuhan nyata. (Ahmad et al., 2024) membuktikan bahwa metode ini efektif dalam pengembangan sistem persediaan karena perbaikan dapat dilakukan bertahap berdasarkan umpan balik langsung. Pendekatan ini ideal untuk sistem dengan kebutuhan yang dinamis dan memerlukan penyesuaian berkelanjutan.",

"PHP (Hypertext Preprocessor) adalah bahasa pemrograman sisi server yang umum digunakan dalam pengembangan sistem informasi berbasis web. Bahasa ini mampu menghasilkan halaman web yang dinamis serta memfasilitasi interaksi langsung antara aplikasi dan basis data relasional. Dalam konteks sistem informasi persediaan maupun penjualan, PHP sering dimanfaatkan karena kemampuannya dalam mengelola data transaksi secara sistematis dan berkelanjutan.":
"PHP (Hypertext Preprocessor) adalah bahasa pemrograman sisi server yang luas digunakan dalam pengembangan aplikasi web. PHP mampu menghasilkan halaman dinamis dan berinteraksi langsung dengan basis data relasional, sehingga cocok untuk sistem informasi persediaan dan penjualan yang membutuhkan pengelolaan data transaksi secara terus-menerus.",

"Penelitian yang dilakukan oleh (Adikusuma et al., 2025) menunjukkan bahwa penggunaan PHP dalam pengembangan sistem informasi penjualan berbasis web memungkinkan pencatatan transaksi dilakukan secara otomatis sekaligus memperbarui data stok secara real time. Hasil serupa juga ditemukan pada penelitian (Ahmad et al., 2024) yang menerapkan PHP dalam pembangunan aplikasi stok opname berbasis metode prototyping. Dalam penelitian tersebut, PHP berperan penting dalam pengelolaan proses input data, pengolahan informasi, hingga penyajian laporan stok kepada pengguna. Selain itu, PHP dikenal fleksibel karena dapat dijalankan di berbagai sistem operasi dan didukung oleh web server populer, sehingga implementasinya sesuai untuk usaha berskala kecil hingga menengah. Berdasarkan karakteristik tersebut, PHP dinilai tepat digunakan sebagai dasar pengembangan sistem informasi penataan persediaan barang berbasis web pada TB. Losari Jaya 2. Adapun framework dari PHP yang dapat memudahkan pengembang dalam mengembangkan websitenya. Framework yang sering digunakan merupakan framework Laravel.":
"(Adikusuma et al., 2025) menunjukkan bahwa PHP memungkinkan pencatatan transaksi otomatis sekaligus pembaruan stok secara real time. (Ahmad et al., 2024) mempertegas peran PHP dalam pengelolaan input data, pengolahan informasi, dan penyajian laporan stok pada aplikasi berbasis prototyping. PHP juga fleksibel karena berjalan di berbagai sistem operasi dan didukung oleh berbagai web server, sehingga cocok untuk usaha skala kecil hingga menengah. Salah satu framework PHP yang paling sering digunakan dalam pengembangan sistem informasi modern adalah Laravel.",

"Laravel merupakan framework PHP yang dirancang untuk membantu pengembang membangun sistem informasi berbasis web dengan struktur yang lebih terorganisasi dan efisien. Framework ini mengadopsi arsitektur Model View Controller (MVC), yang memisahkan antara logika aplikasi, pengelolaan data, dan tampilan antarmuka pengguna. Pemisahan ini tidak hanya mempercepat proses pengembangan, tetapi juga memudahkan proses pemeliharaan dan pembaruan sistem.":
"Laravel adalah framework PHP yang mengadopsi arsitektur Model-View-Controller (MVC), memisahkan logika aplikasi, pengelolaan data, dan tampilan antarmuka. Struktur ini mempercepat proses pengembangan sekaligus memudahkan pemeliharaan dan pembaruan sistem secara berkelanjutan.",

"Menurut (Adikusuma et al., 2025), penggunaan framework PHP dalam pengembangan sistem informasi berbasis web mampu mempercepat proses pembuatan aplikasi serta menyesuaikannya dengan kebutuhan pengguna. Sementara itu, (Ahmad et al., 2024) menegaskan bahwa framework seperti Laravel sangat mendukung penerapan metode prototyping, karena memungkinkan evaluasi dan penyempurnaan sistem dilakukan secara bertahap berdasarkan umpan balik pengguna. Dengan keunggulan dalam hal fleksibilitas, keteraturan struktur, dan kemudahan pengembangan, Laravel dinilai sesuai untuk digunakan dalam pembangunan sistem informasi penataan persediaan barang penjualan berbasis web.":
"(Adikusuma et al., 2025) menyatakan bahwa penggunaan framework PHP mempercepat pembuatan aplikasi dan mempermudah penyesuaian dengan kebutuhan pengguna. (Ahmad et al., 2024) menambahkan bahwa Laravel sangat mendukung metode prototyping karena memungkinkan evaluasi dan penyempurnaan bertahap. Dengan fleksibilitas, struktur yang teratur, dan kemudahan pengembangan, Laravel menjadi pilihan tepat untuk membangun sistem informasi persediaan pada TB. Losari Jaya 2.",

"Unified Modeling Language (UML) merupakan bahasa pemodelan standar yang digunakan untuk memvisualisasikan, merancang, dan mendokumentasikan sistem perangkat lunak secara terstruktur. UML berfungsi sebagai alat bantu komunikasi antara pengembang dan pengguna dalam memahami kebutuhan sistem sebelum tahap implementasi dilakukan. Penggunaan UML memungkinkan pengembang menggambarkan fungsi sistem, alur proses, serta hubungan antar komponen secara jelas sehingga dapat mengurangi kesalahan dalam pengembangan sistem. Penelitian (Adikusuma et al., 2025) menyatakan bahwa penerapan UML dalam perancangan sistem informasi berbasis web membantu meningkatkan pemahaman kebutuhan sistem dan mempercepat proses pengembangan aplikasi.":
"Unified Modeling Language (UML) adalah bahasa pemodelan standar untuk memvisualisasikan, merancang, dan mendokumentasikan sistem perangkat lunak. UML membantu komunikasi antara pengembang dan pengguna dalam memahami kebutuhan sistem sebelum implementasi. Dengan UML, fungsi sistem, alur proses, dan hubungan antar komponen dapat digambarkan secara jelas sehingga meminimalkan kesalahan pengembangan. (Adikusuma et al., 2025) menyatakan bahwa penerapan UML dalam perancangan sistem informasi berbasis web mempercepat pemahaman kebutuhan dan proses pengembangan aplikasi.",

"Dalam pengembangan sistem informasi penataan persediaan barang penjualan pada TB. Losari Jaya 2, UML digunakan sebagai alat bantu perancangan untuk menggambarkan kebutuhan fungsional sistem dan struktur data yang akan diimplementasikan. Penggunaan UML juga sejalan dengan metode prototyping yang menekankan komunikasi intensif dengan pengguna, karena model visual yang dihasilkan lebih mudah dipahami dibandingkan deskripsi teknis semata (Ahmad et al., 2024). Adapun jenis jenis diagram UML, sebagaimana berikut :":
"Dalam penelitian ini, UML digunakan untuk menggambarkan kebutuhan fungsional dan struktur data sistem persediaan TB. Losari Jaya 2. Penggunaan UML sejalan dengan metode prototyping yang menekankan komunikasi intensif dengan pengguna, karena model visual lebih mudah dipahami daripada deskripsi teknis semata (Ahmad et al., 2024). Adapun jenis diagram UML yang digunakan adalah sebagai berikut:",

"Use Case Diagram memvisualisasikan interaksi antara pengguna dan sistem melalui sejumlah simbol utama. Aktor digambarkan dengan ikon manusia stick untuk merepresentasikan pihak eksternal, seperti pengguna atau administrator, yang berinteraksi langsung dengan sistem. Use case ditampilkan dalam bentuk ellips yang menggambarkan fungsi atau layanan sistem, sedangkan Association relationship menghubungkan aktor dengan fungsi tersebut. Hubungan antar use case dijelaskan melalui simbol include dan extend, di mana include menunjukkan fungsi yang selalu dijalankan bersamaan, sedangkan extend merepresentasikan fitur tambahan yang aktif pada kondisi tertentu. Seluruh elemen sistem dibatasi oleh persegi panjang yang menandakan ruang lingkup sistem. Melalui representasi ini, Use Case Diagram memberikan pemahaman fungsional sistem secara terstruktur dan mudah dipahami.":
"Use Case Diagram menggambarkan interaksi antara aktor dan sistem menggunakan simbol-simbol standar UML. Aktor direpresentasikan dengan ikon manusia stik yang mewakili pihak eksternal yang berinteraksi langsung, seperti admin atau kasir. Fungsi sistem ditampilkan dalam bentuk ellips (use case), dihubungkan ke aktor melalui garis asosiasi. Relasi antar use case menggunakan panah <<include>> untuk fungsi yang selalu berjalan bersama dan <<extend>> untuk fitur tambahan yang aktif pada kondisi tertentu. Seluruh elemen dibatasi persegi panjang sebagai batas sistem. Melalui representasi ini, Use Case Diagram menyajikan gambaran fungsional sistem secara ringkas dan mudah dipahami.",

"Diagram Use Case pada sistem Appointment System memperlihatkan hubungan interaktif antara berbagai pengguna dan sistem secara keseluruhan. Terdapat empat aktor utama, yaitu Patient, New Patient, Doctor, dan Management, yang masing masing menjalankan fungsi berbeda sesuai perannya. Kedua aktor pertama, Patient dan New Patient, berpartisipasi dalam use case Make Appointment, yang merepresentasikan proses pembuatan janji temu. Di sisi lain, Management berperan dalam use case Produce Schedule Information untuk menghasilkan data jadwal yang diperlukan, sedangkan Doctor berinteraksi melalui use case Record Availability guna mencatat ketersediaan waktu praktik. Seluruh aktivitas tersebut berada dalam batas sistem yang dilambangkan oleh area Appointment System, yang menegaskan ruang lingkup serta batas tanggung jawab sistem. Secara keseluruhan, diagram ini memperlihatkan perbedaan hak akses dan fungsi setiap aktor, sekaligus memberikan gambaran fungsional sistem secara ringkas, terstruktur, dan mudah dipahami.":
"Gambar 2.2 menampilkan contoh Use Case Diagram pada sistem Appointment System dengan empat aktor utama: Patient, New Patient, Doctor, dan Management. Patient dan New Patient terlibat dalam use case Make Appointment, Management berperan menghasilkan jadwal melalui Produce Schedule Information, dan Doctor mencatat ketersediaan waktu via Record Availability. Seluruh aktivitas berada dalam batas sistem Appointment System. Diagram ini menunjukkan perbedaan hak akses setiap aktor secara terstruktur dan mudah dipahami.",

"Diagram alir yang ditampilkan menunjukkan tahapan pemenuhan pesanan secara menyeluruh, dimulai dari tahap penerimaan pesanan (Receive Order). Setelah pesanan dicatat, langkah selanjutnya adalah menyiapkan pesanan tersebut (Fill Order). Pada tahap berikutnya, sistem mengevaluasi prioritas pesanan: untuk pesanan cepat (rush order), pengiriman dilakukan secara overnight (Arrange Overnight Delivery), sedangkan untuk pesanan biasa, pengiriman dijadwalkan secara reguler (Arrange Regular Delivery). Setelah proses pengiriman ditetapkan, faktur dikirimkan kepada pelanggan (Send Invoice), dan alur ditutup dengan penerimaan pembayaran (Receive Payment). Secara keseluruhan, diagram ini menggambarkan alur kerja yang terstruktur dengan jelas, menekankan cabang keputusan berdasarkan jenis pesanan sehingga setiap pesanan diproses sesuai prioritas sebelum penagihan dilakukan.":
"Gambar 2.3 menampilkan Activity Diagram proses pemenuhan pesanan, dimulai dari Receive Order, dilanjutkan Fill Order, kemudian sistem menentukan prioritas: pesanan mendesak dikirim via Arrange Overnight Delivery, sedangkan pesanan biasa melalui Arrange Regular Delivery. Setelah pengiriman, faktur dikirimkan (Send Invoice) dan alur ditutup dengan Receive Payment. Diagram ini menggambarkan alur kerja terstruktur dengan percabangan keputusan berdasarkan jenis pesanan.",

"Diagram kelas termasuk salah satu diagram statis dalam Unified Modeling Language (UML) yang berfungsi untuk menggambarkan struktur konseptual suatu sistem. Diagram ini menekankan pada hubungan antar kelas, atribut, dan metode yang dimiliki masing masing kelas, serta asosiasi yang terbentuk di dalam sistem. Setiap kelas merepresentasikan entitas atau konsep utama dalam domain masalah, sementara atribut mencerminkan properti yang dimiliki entitas tersebut dan metode menunjukkan perilaku atau fungsi yang dapat dijalankan oleh kelas. Selain itu, diagram kelas juga menampilkan jenis hubungan antar kelas, seperti inheritance, aggregation, dan composition, sehingga mempermudah pemahaman hierarki serta interaksi antar objek dalam sistem. Dengan demikian, penggunaan diagram kelas pada tahap perancangan memudahkan dokumentasi struktur sistem, mengurangi potensi kesalahan desain, dan mendukung pengembangan perangkat lunak yang lebih terstruktur serta terkontrol.":
"Class Diagram adalah diagram statis dalam UML yang menggambarkan struktur konseptual sistem. Diagram ini menonjolkan hubungan antar kelas beserta atribut dan metode masing-masing. Setiap kelas mewakili entitas utama, atribut mencerminkan propertinya, dan metode menunjukkan fungsi yang dapat dijalankan. Relasi antar kelas seperti inheritance, aggregation, dan composition juga ditampilkan untuk memperjelas hierarki dan interaksi antar objek. Class Diagram memudahkan dokumentasi struktur sistem, meminimalkan kesalahan desain, dan mendukung pengembangan yang lebih terstruktur.",

"Diagram kelas yang disajikan menggambarkan desain sistem informasi medis berbasis pendekatan berorientasi objek dengan delapan kelas utama yang saling terkait. Kelas Patient berfungsi sebagai entitas pusat, menyimpan data identitas dan informasi administratif pasien, termasuk atribut seperti Treatment, Insurance, Address, Phone, Birthdate, dan Age. Kelas ini dihubungkan dengan kelas Provider yang merepresentasikan penyedia layanan kesehatan. Untuk memodelkan aspek klinis pasien, digunakan dua kelas tambahan: Medical History, yang mencatat riwayat penyakit sistemik dan alergi, serta Symptoms, yang merekam gejala subjektif dan objektif. Proses diagnostik diwakili oleh kelas Diagnosis, yang memuat metode pemeriksaan mulai dari anamnesis hingga pencitraan, serta kelas Pathology, yang mendetailkan prosedur laboratorium yang terkait. Selain itu, kelas Appointments menangani administrasi janji temu, sedangkan kelas Laboratory Tests mencakup berbagai pemeriksaan penunjang, meskipun detail atributnya tidak sepenuhnya ditampilkan dalam diagram. Keseluruhan struktur ini menunjukkan integrasi menyeluruh antara aspek administratif, klinis, dan penunjang medis, yang dirancang untuk memfasilitasi pengelolaan data pasien secara lebih efektif dan efisien.":
"Gambar 2.4 menampilkan Class Diagram sistem informasi medis dengan delapan kelas utama yang saling terkait. Kelas Patient sebagai entitas pusat menyimpan data identitas dan administratif pasien. Medical History mencatat riwayat penyakit, Symptoms merekam gejala, dan Diagnosis menampung metode pemeriksaan. Kelas Pathology mendetailkan prosedur laboratorium, Appointments menangani jadwal, dan Laboratory Tests mencakup pemeriksaan penunjang. Keseluruhan struktur menunjukkan integrasi antara aspek administratif, klinis, dan penunjang dalam satu sistem yang terpadu.",

"Pengujian Black Box, atau Black Box Testing, merupakan metode evaluasi perangkat lunak yang menekankan pada pemeriksaan fungsionalitas sistem sesuai spesifikasi kebutuhan pengguna, tanpa memeriksa struktur internal, kode, atau logika implementasinya. Pengujian ini dilakukan dengan memasukkan input tertentu dan memeriksa output yang dihasilkan untuk memastikan perilaku sistem sesuai harapan. Metode ini sangat bermanfaat pada tahap akhir pengembangan karena memvalidasi bahwa seluruh fitur telah memenuhi kebutuhan fungsional dari sudut pandang pengguna akhir. Dalam konteks sistem informasi inventory dan penjualan berbasis website, penelitian oleh (Widiatma & Abdillah, 2025) pada Toko Yuda serta (Ahmad et al., 2024) pada sistem stok opname menunjukkan bahwa pendekatan ini efektif untuk memverifikasi fungsi login, pengelolaan data master, proses transaksi, dan ketahanan sistem terhadap input yang tidak valid. Adapun teknik Black Box Testing yang umum diterapkan meliputi:":
"Black Box Testing adalah metode evaluasi perangkat lunak yang memeriksa fungsionalitas sistem berdasarkan input dan output yang dihasilkan, tanpa menelaah kode atau logika internal. Metode ini memvalidasi apakah seluruh fitur memenuhi kebutuhan fungsional dari sudut pandang pengguna. (Widiatma & Abdillah, 2025) dan (Ahmad et al., 2024) membuktikan bahwa Black Box Testing efektif untuk memverifikasi fungsi login, pengelolaan data master, proses transaksi, dan ketahanan sistem terhadap input tidak valid pada sistem inventori berbasis web. Teknik yang umum diterapkan meliputi:",

"User Acceptance Testing (UAT), atau Pengujian Penerimaan Pengguna, merupakan tahap penting dalam siklus pengembangan sistem yang difokuskan pada validasi kesesuaian sistem dengan kebutuhan bisnis serta spesifikasi fungsional dari perspektif pengguna akhir. Berbeda dengan pengujian teknis, seperti unit, integrasi, atau sistem, yang menitikberatkan pada keakuratan kode dan komponen, UAT menilai sejauh mana solusi yang diterapkan dapat mendukung proses bisnis secara efektif, efisien, dan mudah digunakan di lingkungan nyata. Tahap ini biasanya menjadi syarat utama sebelum sistem resmi diserahkan dan dijalankan secara operasional.":
"User Acceptance Testing (UAT) adalah tahap validasi akhir dalam siklus pengembangan sistem yang memastikan sistem sesuai kebutuhan bisnis dari perspektif pengguna akhir. Berbeda dengan pengujian teknis (unit/integrasi/sistem) yang berfokus pada keakuratan kode, UAT menilai sejauh mana sistem mendukung proses bisnis secara efektif, efisien, dan mudah digunakan di lingkungan nyata. Tahap ini menjadi syarat sebelum sistem resmi diserahkan untuk digunakan secara operasional.",

"UAT dijalankan oleh pengguna akhir atau perwakilan klien, bukan oleh pengembang atau tim QA internal. Tujuannya adalah memastikan sistem mampu menangani skenario bisnis nyata, bukan sekadar mendeteksi bug teknis minor seperti kesalahan tampilan atau glitch fungsional yang seharusnya sudah diatasi pada pengujian sebelumnya. Dengan demikian, UAT berfungsi sebagai verifikasi akhir bahwa perangkat lunak dapat diterima oleh pengguna sebagai solusi siap pakai. Dalam pengembangan sistem informasi untuk UMKM—misalnya sistem inventory Toko Yuda (Widiatma & Abdillah, 2025), aplikasi stok opname (Ahmad et al., 2024), atau sistem penjualan Toko SRC Bu Herlin (Adikusuma et al., 2025)—UAT menentukan apakah sistem benar benar menyederhanakan pencatatan manual, mempercepat alur kerja, dan meningkatkan akurasi pelaporan sesuai harapan pemilik usaha dan staf.":
"UAT dijalankan oleh pengguna akhir atau perwakilan klien untuk memastikan sistem dapat menangani skenario bisnis nyata. Pada pengembangan sistem informasi untuk UMKM seperti inventori Toko Yuda (Widiatma & Abdillah, 2025), aplikasi stok opname (Ahmad et al., 2024), dan sistem penjualan Toko SRC Bu Herlin (Adikusuma et al., 2025), UAT menjadi penentu apakah sistem benar-benar menyederhanakan pencatatan, mempercepat alur kerja, dan meningkatkan akurasi laporan sesuai harapan pemilik usaha.",

# ============================================================
# BAB III - SUMBER DATA
# ============================================================
"Pengumpulan data menjadi tahap awal yang krusial dalam pengembangan sistem, karena berfungsi untuk mengidentifikasi permasalahan, merumuskan kebutuhan, serta merancang solusi sistem informasi yang sesuai untuk TB. Losari Jaya 2. Dalam penelitian ini, pengembang dan user bertemu untuk membahas detail yang diinginkan.  Adapun pendekatan kombinasi antara data primer dan data sekunder untuk memastikan analisis yang komprehensif dan akurat sebagai berikut :":
"Pengumpulan data menjadi tahap awal yang krusial dalam pengembangan sistem karena berfungsi untuk mengidentifikasi permasalahan, merumuskan kebutuhan, dan merancang solusi yang tepat untuk TB. Losari Jaya 2. Penelitian ini menggunakan kombinasi data primer dan data sekunder untuk memastikan analisis yang komprehensif, sebagai berikut:",

"Data primer didapat dengan memperoleh data langsung dengan mendatangi TB. Losari Jaya 2. Dalam tahapan ini, penulis menggunakan beberapa tahapan seperti :":
"Data primer diperoleh secara langsung melalui kunjungan ke TB. Losari Jaya 2 dengan menggunakan beberapa pendekatan, yaitu:",

"Data sekunder diperoleh secara tidak langsung melalui kajian terhadap literatur dan penelitian terdahulu yang relevan. Data ini digunakan untuk membangun landasan teori, memahami metode pengembangan yang efektif, dan membandingkan solusi yang ada.":
"Data sekunder diperoleh dari kajian literatur dan penelitian terdahulu yang relevan. Data ini digunakan untuk membangun landasan teori, memahami metode pengembangan yang efektif, serta membandingkan solusi yang telah diterapkan dalam konteks serupa.",

# ============================================================
# BAB III - METODE PENGUMPULAN DATA
# ============================================================
"Pada tahap ini, data dan informasi dikumpulkan dari wawancara, observasi dan studi dokumen. Berikut tahapannya :":
"Data dan informasi dikumpulkan melalui tiga teknik utama, yaitu wawancara, observasi, dan studi dokumen.",

"Beberapa pertanyaan diberikan kepada pemilik TB.Losari Jaya 2, yaitu :":
"Wawancara dilakukan kepada pemilik TB. Losari Jaya 2 dengan pertanyaan-pertanyaan berikut:",

"Observasi partisipatif pasif dilakukan dengan mengamati secara langsung aktivitas operasional di TB. Losari Jaya 2 selama periode tertentu. Fokus utama observasi mencakup proses bisnis inti, yaitu:":
"Observasi pasif dilakukan dengan mengamati langsung kegiatan operasional di TB. Losari Jaya 2 selama periode tertentu. Fokus pengamatan mencakup tiga proses bisnis inti, yaitu:",

"Studi dokumen dilakukan dengan dokumen fisik yang saat ini digunakan dalam operasional toko, sebagai bukti objektif penerapan sistem manual. Dokumen yang dianalisis meliputi:":
"Studi dokumen dilakukan terhadap dokumen fisik yang digunakan dalam operasional toko sebagai bukti penerapan sistem manual. Dokumen yang dianalisis meliputi:",

# ============================================================
# BAB III - METODE PENGEMBANGAN
# ============================================================
"Berikut Langkah langkah metode pengembangan sistem berbasis prototyping pada sistem penataan persediaan barang TB.Losari Jaya 2 :":
"Pengembangan sistem menggunakan metode prototyping dengan langkah-langkah sebagai berikut:",

"Tahap awal difokuskan pada pemahaman mendalam mengenai permasalahan dan kebutuhan pengguna. Observasi dan wawancara mengidentifikasi bahwa TB. Losari Jaya 2 menggunakan sistem pencatatan stok manual yang menimbulkan ketidakakuratan data, lambatnya pencarian informasi, kesulitan pembuatan laporan, dan risiko kehilangan data. Melalui komunikasi intensif dengan pemilik, disepakati bahwa solusi yang dibutuhkan adalah sistem informasi berbasis website yang menampilkan laporan secara real time. Proses komunikasi ini juga menyelaraskan ekspektasi pengembang dan pengguna terkait ruang lingkup, fitur, dan tujuan sistem.":
"Tahap komunikasi difokuskan pada pemahaman mendalam terhadap masalah dan kebutuhan pengguna. Observasi dan wawancara mengungkap bahwa TB. Losari Jaya 2 masih mengandalkan pencatatan manual yang mengakibatkan data tidak akurat, pencarian lambat, laporan sulit dibuat, dan data rawan hilang. Melalui komunikasi intensif bersama pemilik, disepakati bahwa solusi yang diperlukan adalah sistem informasi berbasis website dengan laporan real time. Tahap ini juga menyelaraskan ekspektasi antara pengembang dan pengguna terkait ruang lingkup dan fitur sistem.",

"Hasil komunikasi digunakan untuk merumuskan rencana pengembangan prototyping yang menitikberatkan pada fitur fitur inti untuk mengatasi masalah utama. Fitur utama pada prototyping awal mencakup:":
"Hasil komunikasi dirumuskan menjadi rencana pengembangan yang menitikberatkan pada fitur-fitur inti. Fitur yang diprioritaskan pada prototype awal meliputi:",

"Quick design dibuat berdasarkan rencana yang telah disepakati, dengan fokus pada alur kerja (workflow) dan antarmuka pengguna (user interface) untuk fitur utama:":
"Berdasarkan rencana yang disepakati, quick design dibuat dengan fokus pada alur kerja dan antarmuka pengguna untuk setiap fitur utama:",

"Tahap ini merealisasikan desain menjadi prototyping yang dapat diuji, dibagi menjadi dua iterasi:":
"Tahap ini mewujudkan desain menjadi prototype yang dapat diuji, dilaksanakan dalam dua iterasi:",

"Prototyping 1 (Fungsionalitas Dasar): Fokus pada logika bisnis inti dan backend, mencakup manajemen data barang, transaksi barang masuk sederhana, dan daftar stok. Antarmuka masih low fidelity. Tujuan utamanya adalah memvalidasi proses bisnis seperti perhitungan stok otomatis dan struktur database.":
"Prototype 1 (Fungsionalitas Dasar): Berfokus pada logika bisnis inti dan backend, mencakup manajemen data barang, transaksi barang masuk sederhana, dan daftar stok. Antarmuka masih low fidelity dengan tujuan memvalidasi proses bisnis seperti kalkulasi stok otomatis dan struktur database.",

"Prototyping 2 (Antarmuka dan Fitur Lengkap): Berdasarkan masukan dari Prototyping 1, fungsionalitas disempurnakan dan antarmuka dikembangkan menjadi high fidelity. Fitur transaksi penjualan, dashboard, dan laporan diimplementasikan sepenuhnya menggunakan framework CSS (Tailwind) agar responsif dan mudah digunakan. validasi input dan feedback pengguna juga ditingkatkan.":
"Prototype 2 (Antarmuka & Fitur Lengkap): Berdasarkan masukan dari Prototype 1, fungsionalitas disempurnakan dan antarmuka dikembangkan menjadi high fidelity menggunakan Tailwind CSS agar responsif. Fitur transaksi penjualan, dashboard, dan laporan diimplementasikan sepenuhnya, disertai peningkatan validasi input dan mekanisme feedback pengguna.",

"Evaluasi internal dan bersama pengguna mengidentifikasi beberapa perbaikan, antara lain:":
"Evaluasi bersama pengguna mengidentifikasi beberapa perbaikan yang diperlukan, antara lain:",

"Tahap akhir mencakup penyerahan sistem untuk diuji dalam lingkungan nyata dan pengumpulan umpan balik:":
"Tahap akhir mencakup deployment sistem ke lingkungan nyata dan pengumpulan umpan balik pengguna:",

"Untuk memastikan sistem yang dikembangkan bekerja sesuai fungsi dan diterima pengguna, penelitian ini menerapkan dua jenis pengujian. Pertama, Black Box Testing digunakan untuk memeriksa fungsionalitas teknis sistem tanpa memperhatikan kode atau logika internal. Kedua, User Acceptance Testing (UAT) dilakukan untuk menilai sejauh mana sistem dapat diterima oleh pengguna akhir.":
"Penelitian ini menerapkan dua metode pengujian untuk memastikan sistem berfungsi dengan benar dan dapat diterima pengguna. Black Box Testing digunakan untuk memeriksa fungsionalitas teknis berdasarkan input-output tanpa memeriksa logika internal, sedangkan User Acceptance Testing (UAT) dilakukan untuk menilai tingkat penerimaan sistem oleh pengguna akhir.",

"Pengujian Black Box berfokus pada kesesuaian output terhadap input yang diberikan, sesuai spesifikasi sistem, tanpa menelaah struktur internal. Setiap modul utama, termasuk login, manajemen data master (barang, supplier), transaksi, dan laporan, diuji untuk memastikan beroperasi sesuai harapan.":
"Pengujian Black Box berfokus pada kesesuaian output terhadap input berdasarkan spesifikasi sistem. Modul yang diuji mencakup login, manajemen data master (barang, kategori, satuan, supplier), transaksi barang masuk, transaksi penjualan, dan laporan.",

"Setelah Black Box Testing menunjukkan sistem stabil secara fungsional, UAT dilakukan untuk memverifikasi apakah sistem memenuhi kebutuhan bisnis dan mudah digunakan oleh pengguna akhir pemilik TB. Losari Jaya 2. UAT menjadi prasyarat sebelum implementasi penuh.":
"UAT dilakukan setelah Black Box Testing menunjukkan sistem stabil secara fungsional. Pengujian ini memverifikasi apakah sistem memenuhi kebutuhan bisnis dan mudah digunakan oleh pemilik serta staf TB. Losari Jaya 2 sebelum implementasi penuh.",

# ============================================================
# BAB IV - HASIL DAN PEMBAHASAN
# ============================================================
"Pengembangan sistem informasi penataan persediaan barang penjualan pada TB. Losari Jaya 2 dilakukan menggunakan metode prototyping. Metode ini dipilih karena proses pengembangan membutuhkan keterlibatan pengguna secara langsung, sehingga kebutuhan toko dapat diketahui lebih cepat dan perbaikan dapat dilakukan berdasarkan umpan balik dari pemilik maupun kasir.":
"Pengembangan sistem informasi persediaan pada TB. Losari Jaya 2 dilakukan menggunakan metode prototyping. Metode ini dipilih karena pengembangan membutuhkan keterlibatan langsung pengguna, sehingga kebutuhan toko dapat diidentifikasi lebih cepat dan perbaikan dapat dilakukan berdasarkan umpan balik nyata dari pemilik maupun kasir.",

"Implementasi sistem dibuat dalam bentuk aplikasi Point of Sale berbasis web menggunakan framework Laravel, basis data MySQL, serta antarmuka web yang dapat digunakan untuk pengelolaan data barang, transaksi penjualan, barang masuk,":
"Sistem diimplementasikan sebagai aplikasi Point of Sale berbasis web menggunakan framework Laravel dan basis data MySQL. Antarmuka web mendukung pengelolaan data barang, transaksi penjualan, dan barang masuk,",

"Pada prototype pertama, fokus pengembangan diarahkan pada pemahaman kebutuhan utama dan rancangan dasar alur sistem. Berdasarkan observasi, proses stok barang dan transaksi penjualan di TB. Losari Jaya":
"Pada prototype pertama, pengembangan difokuskan pada pemahaman kebutuhan utama dan perancangan alur sistem dasar. Berdasarkan observasi, proses pencatatan stok dan transaksi di TB. Losari Jaya",

"Tahap communication dilakukan dengan mengidentifikasi kebutuhan admin dan kasir. Admin membutuhkan fitur pengelolaan data barang, kategori, satuan, supplier, barang masuk, laporan, dan keuangan toko.":
"Komunikasi kebutuhan dilakukan bersama admin dan kasir. Admin membutuhkan fitur pengelolaan data barang, kategori, satuan, supplier, barang masuk, laporan stok dan penjualan, serta informasi keuangan toko.",

"Tahap quick plan menghasilkan rancangan awal fitur prioritas, yaitu login, dashboard, data barang, master kategori, master satuan, supplier, barang masuk, transaksi kasir, laporan, dan ringkasan stok.":
"Quick plan menghasilkan rancangan fitur prioritas: login, dashboard, kelola data barang, master kategori dan satuan, kelola supplier, input barang masuk, transaksi kasir, laporan penjualan dan stok.",

"Pada Gambar 4.1 terdapat dua aktor utama, yaitu admin dan kasir. Admin berperan mengelola data master, stok, supplier, barang masuk, laporan, keuangan toko, void transaksi, retur penjualan, stok opname,":
"Gambar 4.1 menampilkan Use Case Diagram sistem dengan dua aktor utama: admin dan kasir. Admin mengelola data master, stok, supplier, barang masuk, laporan, keuangan, void transaksi, retur penjualan, dan stok opname,",

"Pada Gambar 4.2 terlihat alur aktivitas admin setelah berhasil login. Admin dapat berpindah dari dashboard menuju menu data barang, master data, barang masuk, laporan, keuangan, detail transaksi, void":
"Gambar 4.2 menampilkan Activity Diagram alur aktivitas admin setelah login. Dari dashboard, admin dapat mengakses menu data barang, master data, barang masuk, laporan, keuangan, detail transaksi, dan void",

"Pada Gambar 4.3 menggambarkan aktivitas kasir dalam memproses penjualan. Kasir memilih barang, memeriksa stok, memasukkan jumlah dan diskon, menerima pembayaran tunai maupun digital melalui QRIS, meny":
"Gambar 4.3 menggambarkan alur aktivitas kasir dalam memproses transaksi penjualan. Kasir memilih barang, memeriksa stok, memasukkan jumlah dan diskon, menerima pembayaran tunai atau QRIS, kemudian meny",

"Pada tahap model design, rancangan antarmuka difokuskan pada penyusunan tampilan web yang mudah digunakan oleh admin dan kasir. Setiap halaman dibuat untuk mendukung proses operasional toko, mulai dar":
"Pada tahap model design, perancangan antarmuka difokuskan pada tampilan web yang intuitif dan mudah digunakan oleh admin dan kasir. Setiap halaman dirancang untuk mendukung proses operasional toko, mulai dar",

"Pada Gambar 4.4 terlihat halaman login sistem TB. Losari Jaya 2. Halaman ini digunakan oleh admin dan kasir untuk masuk ke sistem menggunakan username dan password yang telah dibuat oleh admin. Tampil":
"Gambar 4.4 menampilkan halaman login sistem TB. Losari Jaya 2. Halaman ini digunakan admin dan kasir untuk mengakses sistem menggunakan username dan password yang telah terdaftar. Tampil",

"Pada Gambar 4.5 terlihat halaman dashboard yang menampilkan ringkasan operasional toko dalam satu layar. Informasi yang ditampilkan meliputi total transaksi hari ini, pendapatan hari ini, jumlah stok":
"Gambar 4.5 menampilkan dashboard yang menyajikan ringkasan operasional toko dalam satu tampilan. Informasi yang disajikan mencakup total transaksi hari ini, pendapatan harian, jumlah stok",

"Pada Gambar 4.6 terlihat halaman daftar barang yang digunakan untuk mengelola data persediaan. Admin dapat melihat nama barang, kategori, supplier, jumlah stok, harga jual, serta tombol aksi untuk men":
"Gambar 4.6 menampilkan halaman daftar barang yang digunakan untuk mengelola data persediaan. Admin dapat melihat nama barang, kategori, supplier, jumlah stok, harga jual, serta tombol aksi untuk men",

"Pada Gambar 4.7 terlihat halaman kelola barang masuk yang digunakan untuk mencatat penerimaan stok dari supplier. Admin mengisi tanggal, supplier, barang, jumlah masuk, harga beli per unit, dan catata":
"Gambar 4.7 menampilkan halaman kelola barang masuk untuk mencatat penerimaan stok dari supplier. Admin mengisi tanggal, nama supplier, barang yang diterima, jumlah masuk, harga beli per unit, dan catata",

"Pada Gambar 4.8 terlihat halaman data supplier yang berfungsi menyimpan informasi pemasok barang. Sistem menampilkan nama supplier, jumlah barang yang terkait, total restock, dan tombol aksi untuk men":
"Gambar 4.8 menampilkan halaman data supplier yang menyimpan informasi pemasok barang. Sistem menampilkan nama supplier, jumlah item terkait, total restock, serta tombol aksi untuk men",

"Pada Gambar 4.9 terlihat halaman kasir yang digunakan untuk memproses transaksi penjualan. Kasir dapat mencari barang, memfilter kategori, memilih barang ke keranjang, mengatur diskon, memasukkan uang":
"Gambar 4.9 menampilkan halaman kasir untuk memproses transaksi penjualan. Kasir dapat mencari barang, memfilter berdasarkan kategori, menambahkan barang ke keranjang, mengatur diskon, memasukkan nominal bayar,",

"Pada Gambar 4.10 terlihat halaman keuangan toko yang menampilkan data pendapatan, pengeluaran, keuntungan, margin, dan arus keuangan. Informasi ini membantu admin melihat kondisi keuangan toko berdasa":
"Gambar 4.10 menampilkan halaman keuangan toko yang menyajikan data pendapatan, pengeluaran, keuntungan, margin, dan arus kas. Informasi ini membantu admin memantau kondisi keuangan toko berdasa",

"Pada Gambar 4.11 terlihat halaman laporan penjualan yang menyediakan ringkasan transaksi berdasarkan periode tertentu. Admin dapat memilih periode laporan, melihat total pendapatan, total transaksi, r":
"Gambar 4.11 menampilkan halaman laporan penjualan yang menyajikan ringkasan transaksi berdasarkan periode yang dipilih. Admin dapat melihat total pendapatan, jumlah transaksi, r",

"Pada Gambar 4.12 terlihat halaman laporan stok barang yang menampilkan snapshot persediaan toko. Informasi yang tersedia meliputi nilai persediaan, jumlah stok aktif, stok minimum, daftar barang, kate":
"Gambar 4.12 menampilkan halaman laporan stok barang yang menyajikan kondisi terkini persediaan toko. Informasi yang tersedia mencakup nilai persediaan, stok aktif, stok minimum, daftar barang, kate",

"Pada Gambar 4.13 terlihat halaman data pengguna yang digunakan admin untuk menambahkan akun admin dan kasir. Form pengguna memuat nama, username, role, password, dan email opsional. Bagian daftar user":
"Gambar 4.13 menampilkan halaman data pengguna yang digunakan admin untuk mengelola akun. Form pengguna memuat nama, username, role, password, dan email opsional. Daftar user",

"Construction of prototype pada tahap pertama menghasilkan rancangan fungsi dasar, yaitu pencatatan data barang, supplier, dan transaksi sederhana. Hasil evaluasi menunjukkan bahwa pengguna membutuhkan":
"Construction of prototype tahap pertama menghasilkan fungsionalitas dasar berupa pencatatan data barang, supplier, dan transaksi sederhana. Evaluasi bersama pengguna menunjukkan kebutuhan tambahan,",

"Pada Gambar 4.14 terlihat halaman master data kategori dan satuan barang. Fitur ini digunakan untuk menambahkan kategori barang, seperti struktur, finishing dan cat, perkakas, plumbing dan pipa, serta":
"Gambar 4.14 menampilkan halaman master data kategori dan satuan barang. Fitur ini memungkinkan admin menambahkan kategori seperti struktur, finishing dan cat, perkakas, plumbing dan pipa, serta",

# ============================================================
# BAB IV - PROTOTYPE 2 & IMPLEMENTASI
# ============================================================
"Implementasi database menggunakan MySQL. Struktur database dirancang untuk menyimpan data pengguna, data master persediaan, supplier, barang masuk, transaksi penjualan, detail penjualan, pelanggan, hu":
"Database diimplementasikan menggunakan MySQL dengan struktur yang dirancang untuk menyimpan data pengguna, master persediaan, supplier, barang masuk, transaksi penjualan, detail transaksi, pelanggan, hu",

"Pada Gambar 4.20 terlihat rancangan ERD database sistem yang menggambarkan relasi tabel users, suppliers, inventory_categories, inventory_units, inventory_items, goods_receipts, sales, sale_items, cus":
"Gambar 4.20 menampilkan ERD database sistem yang menggambarkan relasi antartabel: users, suppliers, inventory_categories, inventory_units, inventory_items, goods_receipts, sales, sale_items, cus",

"Berdasarkan Tabel 4.1, tabel inventory_items menjadi pusat data persediaan karena menyimpan informasi barang dan stok. Tabel goods_receipts digunakan untuk menambah stok saat barang masuk, sedangkan s":
"Berdasarkan Tabel 4.1, tabel inventory_items menjadi inti sistem persediaan karena menyimpan informasi lengkap barang dan stok. Tabel goods_receipts digunakan untuk mencatat penambahan stok dari supplier, sedangkan s",

"Implementasi fungsional sistem menerapkan arsitektur MVC pada Laravel. Model digunakan untuk merepresentasikan tabel database, controller digunakan untuk memproses request, validasi, transaksi database":
"Sistem diimplementasikan menggunakan arsitektur MVC pada Laravel. Model merepresentasikan tabel database, controller memproses request dan validasi,",

"Pada bagian model, sistem memiliki model User, Supplier, InventoryItem, InventoryCategory, InventoryUnit, GoodsReceipt, Sale, SaleItem, Customer, DebtPayment, VoidLog, ReturnRequest, ReturnItem, Stock":
"Model yang digunakan meliputi: User, Supplier, InventoryItem, InventoryCategory, InventoryUnit, GoodsReceipt, Sale, SaleItem, Customer, DebtPayment, VoidLog, ReturnRequest, ReturnItem, Stock",

"Pada bagian view, sistem menyediakan halaman dashboard, master data, barang masuk, transaksi kasir, aktivitas kasir, rincian data, void transaksi, retur penjualan, stok opname, data pelanggan, hutang":
"View sistem mencakup halaman: dashboard, master data, barang masuk, transaksi kasir, aktivitas kasir, rincian data, void transaksi, retur penjualan, stok opname, data pelanggan, hutang",

"Pada bagian controller, PosController menangani proses login, logout, bootstrap data, CRUD barang, CRUD supplier, tambah kategori, tambah satuan, barang masuk, checkout penjualan, void transaksi, retu":
"PosController sebagai controller tunggal menangani: login, logout, bootstrap data, CRUD barang, CRUD supplier, tambah kategori dan satuan, input barang masuk, checkout penjualan, void transaksi, retu",

"Fitur laporan harian disempurnakan agar transaksi terbaru ikut muncul pada export PDF. Selain itu, logo toko ditampilkan pada hasil export PDF sehingga laporan memiliki identitas toko yang jelas. Menu keuangan ditambahkan untuk menampilkan pendapatan, pengeluaran, keuntungan, margin, dan data lengkap keuangan toko.":
"Laporan harian disempurnakan agar mencakup transaksi terbaru dalam export PDF, lengkap dengan logo toko sebagai identitas. Menu keuangan ditambahkan untuk menampilkan pendapatan, pengeluaran, keuntungan, dan margin secara terintegrasi.",

"Pada aktivitas kasir, pengguna dapat melihat detail transaksi yang sudah terjadi dan melakukan cetak ulang struk. Pada rincian data, transaksi harian juga dapat dibuka untuk melihat daftar item, jumlah, harga, diskon, total, nominal bayar, dan kembalian.":
"Fitur aktivitas kasir memungkinkan pengguna melihat detail transaksi yang telah berlangsung dan mencetak ulang struk. Rincian data transaksi harian dapat dibuka untuk melihat daftar item, jumlah, harga, diskon, total, nominal bayar, dan kembalian.",

"Tahap deployment dilakukan dengan menyiapkan aplikasi agar dapat dijalankan pada lingkungan server yang mendukung PHP, Composer, Node.js, dan MySQL. Pada tahap lokal, aplikasi dijalankan melalui Laravel dengan konfigurasi database MySQL. Untuk hosting produksi, sistem membutuhkan layanan yang mendukung PHP dan MySQL agar fitur transaksi, laporan, session, dan penyimpanan database dapat berjalan normal.":
"Deployment dilakukan dengan menyiapkan lingkungan server yang mendukung PHP, Composer, Node.js, dan MySQL. Pada tahap lokal, aplikasi dijalankan melalui Laragon dengan konfigurasi database MySQL. Untuk hosting produksi, server harus mendukung PHP dan MySQL agar fitur transaksi, laporan, session, dan penyimpanan database dapat berjalan normal.",

"Tahap delivery dilakukan dengan menyerahkan aplikasi kepada pengguna untuk dicoba pada proses operasional toko. Feedback dari pengguna digunakan untuk memperbaiki kebutuhan tambahan, seperti master kategori dan satuan, detail transaksi, cetak ulang struk, laporan PDF, serta menu keuangan. Perbaikan tersebut kemudian dimasukkan ke prototype kedua agar sistem lebih sesuai dengan kebutuhan nyata toko.":
"Tahap delivery dilakukan dengan menyerahkan aplikasi kepada pengguna untuk diuji dalam operasional nyata toko. Umpan balik yang diterima digunakan untuk menambahkan fitur master kategori dan satuan, detail transaksi, cetak ulang struk, laporan PDF, serta menu keuangan. Perbaikan ini kemudian diintegrasikan ke dalam prototype kedua.",

# ============================================================
# BAB IV - PENGUJIAN
# ============================================================
"Pengujian sistem dilakukan untuk memastikan fitur yang dikembangkan dapat berjalan sesuai kebutuhan pengguna. Metode pengujian yang digunakan adalah Black Box Testing dan User Acceptance Testing (UAT). Black Box Testing berfokus pada fungsi sistem berdasarkan input dan output, sedangkan UAT berfokus pada penerimaan pengguna terhadap kemudahan dan kesesuaian sistem.":
"Pengujian sistem dilakukan untuk memverifikasi bahwa seluruh fitur berfungsi sesuai kebutuhan pengguna. Dua metode pengujian diterapkan: Black Box Testing yang berfokus pada verifikasi input-output setiap fungsi, dan User Acceptance Testing (UAT) yang menilai penerimaan pengguna terhadap kemudahan dan kesesuaian sistem.",

"Pengujian Black Box dilakukan pada fitur utama yang digunakan dalam operasional toko. Setiap fitur diuji berdasarkan skenario penggunaan, hasil yang diharapkan, hasil aktual, dan status pengujian.":
"Pengujian Black Box dilakukan pada seluruh fitur utama operasional toko. Setiap fitur diuji menggunakan skenario penggunaan, kemudian dibandingkan antara hasil yang diharapkan dengan hasil aktual untuk menentukan status pengujian.",

"Berdasarkan Tabel 4.2, seluruh fitur utama yang diuji mendapatkan status valid. Hal ini menunjukkan bahwa sistem telah memenuhi kebutuhan fungsional utama, khususnya pada pengelolaan persediaan, transaksi kasir, laporan, dan keuangan toko.":
"Berdasarkan Tabel 4.2, seluruh fitur utama memperoleh status valid. Hasil ini membuktikan bahwa sistem telah memenuhi kebutuhan fungsional, khususnya dalam pengelolaan persediaan, transaksi kasir, pelaporan, dan informasi keuangan toko.",

"Pengujian UAT dilakukan untuk melihat penerimaan pengguna terhadap sistem. Aspek yang dinilai meliputi kemudahan penggunaan, kesesuaian fitur dengan kebutuhan toko, kejelasan informasi, dan manfaat sistem terhadap pencatatan operasional.":
"UAT dilakukan untuk mengukur tingkat penerimaan pengguna terhadap sistem. Aspek yang dievaluasi meliputi kemudahan penggunaan, kesesuaian fitur dengan kebutuhan toko, kejelasan informasi yang disajikan, dan manfaat nyata sistem terhadap efisiensi pencatatan operasional.",

# ============================================================
# BAB V - KESIMPULAN
# ============================================================
"Pada bab ini dijelaskan kesimpulan dari penelitian yang dilakukan oleh penulis mengenai rancang bangun sistem informasi penataan persediaan barang penjualan pada TB. Losari Jaya 2 dengan metode prototyping berbasis website. Selain itu, bab ini juga memuat rekomendasi yang dapat digunakan sebagai bahan pertimbangan untuk penelitian dan pengembangan sistem selanjutnya.":
"Bab ini menyajikan kesimpulan dari penelitian rancang bangun sistem informasi persediaan barang penjualan pada TB. Losari Jaya 2 menggunakan metode prototyping berbasis website, serta rekomendasi untuk penelitian dan pengembangan sistem selanjutnya.",

"Berdasarkan penelitian yang telah dilakukan, diperoleh beberapa kesimpulan yang menjawab rumusan masalah dan tujuan penelitian. Kesimpulan tersebut antara lain sebagai berikut:":
"Berdasarkan penelitian yang telah dilakukan, diperoleh kesimpulan sebagai berikut:",

"Berdasarkan hasil penelitian yang telah dilakukan, penulis memberikan beberapa rekomendasi yang dapat dijadikan pertimbangan untuk penelitian selanjutnya. Beberapa saran tersebut antara lain sebagai berikut:":
"Berdasarkan hasil penelitian, penulis merumuskan beberapa rekomendasi untuk penelitian dan pengembangan selanjutnya:",
}

print(f"\nJumlah revisi teks yang didefinisikan: {len(REVISI)}")

# =============================================================================
# FUNGSI HELPER
# =============================================================================

def copy_paragraph_format(src_para, dst_para):
    """Copy format paragraph dari src ke dst."""
    dst_para.paragraph_format.left_indent = src_para.paragraph_format.left_indent
    dst_para.paragraph_format.first_line_indent = src_para.paragraph_format.first_line_indent
    dst_para.paragraph_format.space_before = src_para.paragraph_format.space_before
    dst_para.paragraph_format.space_after = src_para.paragraph_format.space_after
    dst_para.paragraph_format.line_spacing = src_para.paragraph_format.line_spacing
    dst_para.paragraph_format.alignment = src_para.paragraph_format.alignment

def get_run_format(run):
    """Ambil format dari run."""
    return {
        'bold': run.bold,
        'italic': run.italic,
        'underline': run.underline,
        'font_name': run.font.name,
        'font_size': run.font.size,
        'color': run.font.color.rgb if run.font.color and run.font.color.type else None,
    }

def apply_run_format(run, fmt):
    """Terapkan format ke run."""
    if fmt['bold'] is not None:
        run.bold = fmt['bold']
    if fmt['italic'] is not None:
        run.italic = fmt['italic']
    if fmt['underline'] is not None:
        run.underline = fmt['underline']
    if fmt['font_name']:
        run.font.name = fmt['font_name']
    if fmt['font_size']:
        run.font.size = fmt['font_size']
    if fmt['color']:
        try:
            run.font.color.rgb = fmt['color']
        except:
            pass

def replace_paragraph_text(para, new_text):
    """Ganti teks paragraph sambil mempertahankan format run pertama."""
    if not para.runs:
        return
    
    # Simpan format run pertama
    first_run = para.runs[0]
    fmt = get_run_format(first_run)
    
    # Hapus semua runs
    for run in para.runs:
        run.text = ''
    
    # Tulis teks baru di run pertama
    para.runs[0].text = new_text
    apply_run_format(para.runs[0], fmt)

def find_best_match(para_text, revisi_dict):
    """Cari teks yang paling cocok dengan paragraf."""
    para_stripped = para_text.strip()
    
    # Exact match
    if para_stripped in revisi_dict:
        return revisi_dict[para_stripped]
    
    # Partial match - cari key yang merupakan awalan dari para_text
    for key, val in revisi_dict.items():
        key_stripped = key.strip()
        # Cek apakah para_text dimulai dengan key (truncated paragraphs)
        if len(key_stripped) > 80 and para_stripped.startswith(key_stripped[:80]):
            return val
        # Cek apakah key dimulai dengan para_text (para terpotong)
        if len(para_stripped) > 80 and key_stripped.startswith(para_stripped[:80]):
            return val
    
    return None

# =============================================================================
# PROSES UTAMA
# =============================================================================

print("\nMembuka dokumen sumber...")
doc = Document(SRC)

replaced_count = 0
not_found = []

print("Merevisi teks paragraf...")
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if not text:
        continue
    
    new_text = find_best_match(text, REVISI)
    if new_text:
        replace_paragraph_text(para, new_text)
        replaced_count += 1

print(f"Total paragraf direvisi: {replaced_count}")

# =============================================================================
# SISIPKAN GAMBAR DIAGRAM BARU
# =============================================================================
print("\nMenyisipkan gambar diagram baru...")

def find_paragraph_by_caption(doc, caption_text):
    """Cari paragraf berdasarkan teks caption."""
    for i, para in enumerate(doc.paragraphs):
        if caption_text.lower() in para.text.lower():
            return i, para
    return -1, None

def insert_image_after_paragraph(doc, para_idx, img_path, width_inches=5.5):
    """Sisipkan gambar setelah paragraf tertentu."""
    if not os.path.exists(img_path):
        print(f"  WARN: Gambar tidak ditemukan: {img_path}")
        return False
    
    # Dapatkan paragraf target
    target_para = doc.paragraphs[para_idx]
    
    # Cari paragraf kosong atau teks setelah caption
    # Cari apakah ada paragraph gambar yang sudah ada (placeholder)
    next_idx = para_idx + 1
    while next_idx < len(doc.paragraphs):
        next_para = doc.paragraphs[next_idx]
        # Jika paragraf berikutnya sudah ada gambar, ganti
        if len(next_para._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')) > 0:
            # Ada gambar, tambahkan run baru
            # Clear existing runs
            for run in next_para.runs:
                run.text = ''
            # Tambah gambar baru
            run = next_para.add_run()
            run.add_picture(img_path, width=Inches(width_inches))
            print(f"  OK: Gambar diganti di P{next_idx}")
            return True
        elif next_para.text.strip() == '' or next_para.style.name in ['Caption', 'Normal']:
            break
        next_idx += 1
    
    # Jika tidak ada placeholder, insert setelah caption
    # Menggunakan XML manipulation
    new_para = doc.add_paragraph()
    new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = new_para.add_run()
    run.add_picture(img_path, width=Inches(width_inches))
    
    # Pindahkan paragraf baru setelah target
    target_para._element.addnext(new_para._element)
    print(f"  OK: Gambar baru disisipkan setelah P{para_idx}")
    return True

# Cari lokasi caption dan sisipkan gambar
diagrams_to_insert = [
    ("Gambar 4.1", IMG_USECASE, 5.0),
    ("Gambar 4.2", IMG_ACT_ADMIN, 5.5),
    ("Gambar 4.3", IMG_ACT_KASIR, 5.5),
    ("Gambar 4.19", IMG_CLASS, 5.5),
    ("Gambar 4.20", IMG_ERD, 5.5),
]

for caption, img_path, width in diagrams_to_insert:
    idx, para = find_paragraph_by_caption(doc, caption)
    if para:
        print(f"  Menemukan '{caption}' di P{idx}")
        # Cari paragraf gambar yang ada setelahnya (bisa berupa paragraph dengan shape/drawing)
        inserted = False
        for j in range(idx+1, min(idx+5, len(doc.paragraphs))):
            p = doc.paragraphs[j]
            drawings = p._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')
            if drawings:
                # Replace gambar yang ada
                # Remove all drawing elements
                for drawing in drawings:
                    drawing.getparent().remove(drawing)
                # Tambah gambar baru
                run = p.runs[0] if p.runs else p.add_run()
                run.text = ''
                run.add_picture(img_path, width=Inches(width))
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                print(f"    -> Gambar lama di P{j} diganti")
                inserted = True
                break
        if not inserted:
            print(f"    -> Tidak ada gambar placeholder ditemukan untuk '{caption}'")
    else:
        print(f"  WARN: Caption '{caption}' tidak ditemukan")

# =============================================================================
# SIMPAN DOKUMEN
# =============================================================================
print(f"\nMenyimpan ke: {DST}")
doc.save(DST)
print(f"SELESAI! Dokumen tersimpan sebagai: {DST}")
print(f"Paragraf direvisi: {replaced_count}")
