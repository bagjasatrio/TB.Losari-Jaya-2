from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "TA2_14694_REVISI.docx"


def main() -> None:
    document = Document(SOURCE)
    for index, paragraph in enumerate(document.paragraphs):
        drawings = paragraph._p.xpath(".//w:drawing")
        if not drawings:
            continue

        refs = paragraph._p.xpath(".//a:blip/@r:embed")
        names = []
        for relationship_id in refs:
            part = document.part.related_parts[relationship_id]
            names.append(Path(part.partname).name)

        next_text = ""
        for candidate in document.paragraphs[index + 1 : index + 4]:
            if candidate.text.strip():
                next_text = candidate.text.strip()
                break

        extents = []
        for extent in paragraph._p.xpath(".//wp:extent"):
            extents.append(
                (
                    extent.get(qn("wp:cx")),
                    extent.get(qn("wp:cy")),
                )
            )
        print(f"P{index}: images={names}, extents={extents}, next={next_text[:100]}")

    with ZipFile(SOURCE) as archive:
        media = sorted(name for name in archive.namelist() if name.startswith("word/media/"))
        print(f"\nEmbedded media: {len(media)}")
        for name in media:
            info = archive.getinfo(name)
            print(f"{name}\t{info.file_size}")


if __name__ == "__main__":
    main()
