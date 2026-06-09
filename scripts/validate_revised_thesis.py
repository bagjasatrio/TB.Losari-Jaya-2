import re
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "TA2_14694_REVISI.docx"
OUTPUT = ROOT / "TA2_14694_REVISI_BARU.docx"


def citations(document):
    text = "\n".join(p.text for p in document.paragraphs)
    return sorted(set(re.findall(r"\([^()]*?(?:19|20)\d{2}[^()]*?\)", text)))


def main():
    source = Document(SOURCE)
    output = Document(OUTPUT)
    assert len(source.tables) == len(output.tables)

    # Cover, approval, examination, originality, and publication pages must remain unchanged.
    # The automatic tables of contents/figures begin at P79 and are expected to update page numbers.
    for index in range(79):
        assert source.paragraphs[index].text == output.paragraphs[index].text, f"Front matter changed at P{index}"
        assert source.paragraphs[index].style.name == output.paragraphs[index].style.name, f"Front style changed at P{index}"

    for table_index in range(len(source.tables)):
        source_text = [[cell.text for cell in row.cells] for row in source.tables[table_index].rows]
        output_text = [[cell.text for cell in row.cells] for row in output.tables[table_index].rows]
        assert source_text == output_text, f"Table {table_index} changed"

    source_bibliography_start = next(i for i, p in enumerate(source.paragraphs) if p.text.strip() == "DAFTAR PUSTAKA")
    output_bibliography_start = next(i for i, p in enumerate(output.paragraphs) if p.text.strip() == "DAFTAR PUSTAKA")
    source_bibliography = [p.text for p in source.paragraphs[source_bibliography_start:]]
    output_bibliography = [p.text for p in output.paragraphs[output_bibliography_start:]]
    assert source_bibliography == output_bibliography, "Bibliography changed"

    source_citations = citations(source)
    output_citations = citations(output)
    missing = sorted(set(source_citations) - set(output_citations))
    assert not missing, f"Citations removed: {missing}"

    captions = [p.text for p in output.paragraphs if p.style.name == "Caption"]
    assert len(captions) == len([p.text for p in source.paragraphs if p.style.name == "Caption"])
    assert captions == [p.text for p in source.paragraphs if p.style.name == "Caption"]

    diagram_captions = [
        "Gambar 4.1 Use Case Diagram Sistem TB. Losari Jaya 2",
        "Gambar 4.2 Activity Diagram Admin",
        "Gambar 4.3 Activity Diagram Kasir",
        "Gambar 4.18 Sequence Diagram Transaksi Penjualan",
        "Gambar 4.19 Class Diagram Sistem",
        "Gambar 4.20 ERD Database Sistem",
    ]
    for caption in diagram_captions:
        index = next(i for i, p in enumerate(output.paragraphs) if p.text.strip() == caption)
        assert output.paragraphs[index - 1]._p.xpath(".//w:drawing"), f"Missing diagram before {caption}"

    print("Validation passed")
    print(f"Paragraphs: {len(output.paragraphs)} (automatic fields updated)")
    print(f"Tables: {len(output.tables)}")
    print(f"Citations retained: {len(output_citations)} unique citation groups")
    print(f"Captions retained: {len(captions)}")


if __name__ == "__main__":
    main()
