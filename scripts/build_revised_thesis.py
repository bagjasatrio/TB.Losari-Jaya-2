from pathlib import Path
from shutil import copy2

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "TA2_14694_REVISI.docx"
OUTPUT = ROOT / "TA2_14694_REVISI_BARU.docx"
DIAGRAMS = ROOT / "diagram_revisi"


IMAGE_REPLACEMENTS = {
    354: "usecase_pos.png",
    357: "activity_admin.png",
    360: "activity_cashier.png",
    412: "sequence_transaksi.png",
    415: "classdiagram_pos.png",
    420: "erddiagram_pos.png",
}


TEXT_REPLACEMENTS = {
    122: (
        "TB. Losari Jaya 2 memiliki perputaran persediaan yang tinggi, tetapi pencatatan barang masuk, "
        "barang keluar, dan transaksi penjualan masih dilakukan secara manual. Kondisi tersebut dapat "
        "menimbulkan selisih stok, memperlambat pencarian informasi, serta meningkatkan risiko kehilangan "
        "data. Penelitian ini bertujuan membangun sistem informasi persediaan dan penjualan berbasis website "
        "yang membantu toko mengelola stok dan transaksi secara lebih akurat. Sistem dikembangkan menggunakan "
        "metode Prototyping melalui tahap komunikasi kebutuhan, perancangan awal, pembangunan prototype, "
        "evaluasi pengguna, dan penyempurnaan. Implementasi menggunakan framework Laravel dan basis data MySQL. "
        "Hasil pengujian Black Box dan User Acceptance Testing menunjukkan bahwa fungsi utama sistem berjalan "
        "sesuai kebutuhan serta membantu admin dan kasir mencatat transaksi, memperbarui stok, dan menyusun laporan."
    ),
    128: (
        "Kemajuan teknologi informasi membantu pelaku usaha mengelola data operasional secara lebih cepat dan "
        "akurat. Pada usaha perdagangan, persediaan berhubungan langsung dengan kelancaran penjualan dan kualitas "
        "pelayanan. Pencatatan manual sering menimbulkan perbedaan antara catatan dan stok fisik, keterlambatan "
        "informasi, serta kesulitan penyusunan laporan (Gultom, 2022). Karena itu, sistem informasi terintegrasi "
        "diperlukan untuk mendukung pengendalian persediaan."
    ),
    129: (
        "Hasil observasi di TB. Losari Jaya 2 menunjukkan bahwa stok dan transaksi masih dicatat menggunakan buku "
        "tulis. Cara tersebut menyulitkan pencarian riwayat, meningkatkan risiko kehilangan informasi, dan dapat "
        "menyebabkan ketidaksesuaian stok. Permasalahan serupa juga ditemukan pada usaha dagang lain yang belum "
        "menggunakan sistem digital (Widiatma & Abdillah, 2025)."
    ),
    130: (
        "Penelitian terdahulu menunjukkan bahwa sistem persediaan berbasis web dapat mempercepat pencarian data, "
        "mengurangi kesalahan input, dan meningkatkan ketepatan laporan (Effendi & Noviansyah, 2018). Metode "
        "prototyping juga mendukung keterlibatan pengguna sejak tahap awal sehingga rancangan dapat disesuaikan "
        "dengan kebutuhan operasional (Ahmad et al., 2024)."
    ),
    131: (
        "Integrasi penjualan dan persediaan membuat stok dapat diperbarui secara otomatis setelah transaksi. "
        "Adikusuma et al. (2025) menunjukkan bahwa sistem penjualan berbasis web dengan metode prototyping mampu "
        "mengurangi kesalahan pencatatan dan mempercepat penyusunan informasi penjualan."
    ),
    132: (
        "Berdasarkan kondisi tersebut, penelitian ini menerapkan metode prototyping untuk membangun sistem "
        "persediaan dan penjualan yang sesuai dengan proses kerja TB. Losari Jaya 2. Sistem diharapkan membantu "
        "toko mencatat transaksi secara terstruktur, memantau stok, dan memperoleh laporan dengan lebih cepat."
    ),
    135: "Berdasarkan latar belakang tersebut, rumusan masalah penelitian ini adalah sebagai berikut:",
    136: "Bagaimana kondisi pengelolaan persediaan dan transaksi yang masih dilakukan secara manual di TB. Losari Jaya 2?",
    137: "Bagaimana merancang dan membangun sistem informasi persediaan berbasis website menggunakan metode prototyping?",
    140: "Agar penelitian tetap terarah, ruang lingkup penelitian dibatasi sebagai berikut:",
    141: "Penelitian berfokus pada perancangan dan pengembangan aplikasi web untuk mengelola persediaan dan penjualan di TB. Losari Jaya 2.",
    145: "Tujuan penelitian ini adalah sebagai berikut:",
    150: "Penelitian ini diharapkan memberikan manfaat teoritis dan praktis sebagai berikut:",
    168: (
        "Pengelolaan persediaan menentukan ketersediaan barang dan kelancaran transaksi pada usaha dagang. "
        "Pencatatan manual dapat menimbulkan selisih stok, keterlambatan informasi, dan kesulitan penyusunan "
        "laporan. Gultom (2022) menunjukkan bahwa sistem informasi berbasis web dapat meningkatkan keteraturan "
        "dan ketepatan data persediaan."
    ),
    169: (
        "Sistem persediaan juga membantu mempercepat pencarian data dan penyusunan laporan berkala "
        "(Effendi & Noviansyah, 2018). Selain itu, sistem inventori berbasis web dapat mengurangi risiko "
        "kehilangan data dan meningkatkan ketepatan pencatatan arus barang (Widiatma & Abdillah, 2025)."
    ),
    170: (
        "Metode pengembangan turut menentukan kesesuaian sistem dengan kebutuhan pengguna. Prototyping "
        "memungkinkan pengguna terlibat dalam perancangan dan evaluasi sistem. Ahmad et al. (2024) menunjukkan "
        "bahwa metode ini dapat menghasilkan aplikasi stok opname yang lebih responsif, sedangkan Adikusuma "
        "et al. (2025) menunjukkan manfaatnya dalam mengurangi kesalahan pencatatan penjualan. Oleh karena itu, "
        "metode prototyping dipilih untuk pengembangan sistem di TB. Losari Jaya 2."
    ),
    203: (
        "Pada penelitian ini, UML digunakan untuk menggambarkan kebutuhan fungsional, alur aktivitas, dan struktur "
        "sistem. Model visual memudahkan komunikasi dengan pengguna selama proses prototyping karena rancangan "
        "dapat dipahami dan dievaluasi sebelum diimplementasikan (Ahmad et al., 2024). Diagram UML yang digunakan "
        "dijelaskan sebagai berikut:"
    ),
    208: (
        "Use Case Diagram menggambarkan hubungan antara aktor dan fungsi sistem. Aktor menunjukkan pihak yang "
        "menggunakan sistem, use case menunjukkan layanan yang tersedia, sedangkan garis asosiasi menunjukkan "
        "interaksi keduanya. Hubungan include digunakan untuk proses yang selalu dipanggil, sementara extend "
        "menunjukkan fungsi tambahan pada kondisi tertentu."
    ),
    225: (
        "Activity Diagram menggambarkan urutan aktivitas, keputusan, dan perpindahan proses dalam sistem. Diagram "
        "ini membantu menjelaskan alur kerja pengguna dan respons sistem secara ringkas."
    ),
    233: (
        "Class Diagram menggambarkan struktur statis sistem melalui kelas, atribut, metode, dan hubungan "
        "antarkelas. Diagram ini membantu pengembang memahami tanggung jawab setiap objek serta keterkaitan data "
        "sebelum proses implementasi."
    ),
    346: (
        "Sistem diimplementasikan sebagai aplikasi Point of Sale berbasis web menggunakan Laravel dan MySQL. "
        "Aplikasi mendukung pengelolaan barang, barang masuk, transaksi penjualan, pelanggan, laporan, dan "
        "informasi keuangan toko."
    ),
    356: (
        "Gambar 4.1 menunjukkan dua aktor, yaitu admin dan kasir. Admin mengelola data master, persediaan, "
        "pengguna, pelanggan, hutang, void, retur, stok opname, laporan, dan keuangan. Kasir memproses penjualan, "
        "menerima pembayaran, melihat detail transaksi, serta mencetak struk."
    ),
    359: (
        "Gambar 4.2 menunjukkan alur kerja admin. Setelah membuka dashboard, admin memilih fitur administrasi, "
        "mengisi atau memperbarui data, lalu sistem melakukan validasi sebelum menyimpan perubahan. Jika data "
        "valid, database diperbarui dan sistem mengirimkan state terbaru ke dashboard."
    ),
    362: (
        "Gambar 4.3 menunjukkan alur transaksi kasir. Kasir memilih barang, sistem memeriksa stok dan menghitung "
        "total, kemudian kasir memilih pelanggan serta metode pembayaran. Setelah validasi berhasil, transaksi "
        "disimpan, stok dikurangi, dan struk dapat dicetak."
    ),
    410: (
        "Prototype kedua dikembangkan dari hasil evaluasi prototype pertama. Penyempurnaan meliputi master "
        "kategori dan satuan, detail transaksi, cetak ulang struk, laporan PDF, menu keuangan, void transaksi, "
        "retur, stok opname, pelanggan, serta pencatatan hutang."
    ),
    411: (
        "Pada prototype kedua, pemodelan difokuskan pada alur transaksi dan struktur data aplikasi. Sequence "
        "diagram menjelaskan urutan checkout, sedangkan class diagram dan ERD menunjukkan hubungan model serta "
        "tabel yang mendukung proses operasional."
    ),
    414: (
        "Gambar 4.18 menunjukkan urutan transaksi penjualan. Kasir memilih barang dan mengirim permintaan checkout. "
        "PosController memvalidasi stok dan pembayaran, menyimpan penjualan beserta itemnya dalam transaksi "
        "database, mengurangi stok, lalu mengirim state terbaru sebelum struk dicetak."
    ),
    417: (
        "Gambar 4.19 menunjukkan kelas utama sistem dan relasinya. InventoryItem terhubung dengan supplier, "
        "kategori, satuan, barang masuk, item penjualan, dan item stok opname. Sale terhubung dengan kasir, "
        "pelanggan, detail penjualan, pembayaran hutang, void, dan retur. Struktur tersebut menjaga pencatatan "
        "stok dan transaksi tetap saling terhubung."
    ),
    419: (
        "Database MySQL menyimpan data pengguna, master persediaan, supplier, barang masuk, penjualan, pelanggan, "
        "hutang, void, retur, dan stok opname. Relasi antartabel dirancang agar perubahan transaksi dan stok dapat "
        "dicatat secara konsisten."
    ),
    422: (
        "Gambar 4.20 menunjukkan hubungan tabel utama sistem. Tabel inventory_items menjadi pusat data persediaan "
        "dan terhubung dengan barang masuk, detail penjualan, serta stok opname. Tabel sales terhubung dengan "
        "pengguna, pelanggan, item penjualan, pembayaran hutang, void, dan retur."
    ),
    427: (
        "Implementasi fungsional menggunakan pola MVC Laravel. Model merepresentasikan data, PosController "
        "menangani request dan validasi, sedangkan Blade dan JavaScript membentuk antarmuka pengguna."
    ),
    430: (
        "PosController menangani autentikasi, bootstrap data, pengelolaan master, barang masuk, checkout, void, "
        "retur, stok opname, pelanggan, pembayaran hutang, reset data demo, dan ekspor laporan. Proses yang "
        "memengaruhi transaksi dan stok dijalankan dalam transaksi database agar perubahan tetap konsisten."
    ),
    433: (
        "Deployment dilakukan pada lingkungan yang mendukung PHP, Composer, Node.js, dan MySQL. Pada tahap lokal, "
        "aplikasi dijalankan menggunakan konfigurasi Laravel dan database MySQL sebelum disiapkan untuk server produksi."
    ),
    435: (
        "Aplikasi kemudian diberikan kepada pengguna untuk diuji pada proses operasional toko. Masukan pengguna "
        "digunakan untuk menyempurnakan fitur master data, detail transaksi, cetak ulang struk, laporan PDF, "
        "keuangan, dan pengendalian stok."
    ),
    437: (
        "Pengujian dilakukan menggunakan Black Box Testing dan User Acceptance Testing (UAT). Black Box Testing "
        "memeriksa kesesuaian fungsi berdasarkan input dan output, sedangkan UAT menilai penerimaan pengguna "
        "terhadap kemudahan dan manfaat sistem."
    ),
    449: (
        "Bab ini menyajikan kesimpulan penelitian mengenai pengembangan sistem persediaan dan penjualan berbasis "
        "website di TB. Losari Jaya 2 serta rekomendasi untuk pengembangan selanjutnya."
    ),
    451: "Berdasarkan hasil perancangan, implementasi, dan pengujian, diperoleh kesimpulan sebagai berikut:",
}


def clear_paragraph(paragraph):
    paragraph_element = paragraph._element
    for child in list(paragraph_element):
        if child.tag.endswith("}pPr"):
            continue
        paragraph_element.remove(child)


def main():
    copy2(SOURCE, OUTPUT)
    document = Document(OUTPUT)

    for index, text in TEXT_REPLACEMENTS.items():
        document.paragraphs[index].text = text

    available_width = document.sections[0].page_width - document.sections[0].left_margin - document.sections[0].right_margin
    for index, filename in IMAGE_REPLACEMENTS.items():
        paragraph = document.paragraphs[index]
        clear_paragraph(paragraph)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run().add_picture(str(DIAGRAMS / filename), width=available_width)

    document.save(OUTPUT)
    print(f"Created {OUTPUT}")


if __name__ == "__main__":
    main()
