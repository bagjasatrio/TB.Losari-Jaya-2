
import docx
import json

doc = docx.Document(r'c:\Users\ASUS\Documents\Codex\Losarijaya2\TA2_14694_REVISI.docx')

with open(r'c:\Users\ASUS\Documents\Codex\Losarijaya2\doc_content.txt', 'w', encoding='utf-8') as f:
    f.write(f'TOTAL PARAGRAPHS: {len(doc.paragraphs)}\n')
    f.write(f'TOTAL TABLES: {len(doc.tables)}\n\n')
    
    for i, p in enumerate(doc.paragraphs):
        text = p.text
        f.write(f'[P{i}][{p.style.name}] {text}\n')
    
    f.write('\n\n=== TABLES ===\n')
    for t_idx, table in enumerate(doc.tables):
        f.write(f'\n--- Table {t_idx} ---\n')
        for row in table.rows:
            cells = [cell.text for cell in row.cells]
            f.write(' | '.join(cells) + '\n')

print('Done! Saved to doc_content.txt')
